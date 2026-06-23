from __future__ import annotations

from copy import deepcopy
import time
import unicodedata
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Callable

try:
    from zoneinfo import ZoneInfo
except ModuleNotFoundError:
    def ZoneInfo(name: str) -> timezone:
        if name == "Asia/Shanghai":
            return timezone(timedelta(hours=8))
        return timezone.utc

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate, RichText

from .config import (
    CONTRACTS_DIR,
    TemplateConfig,
    TemplateTypography,
    ensure_storage,
    get_template_typography,
    template_docx_path,
)
from .docx_typography import normalize_run_if_body

LEFT_ALIGNED_HEADER_TABLE_COUNT = 2
UNDERLINE_BLANK = "        "
SMALL_CELL_LIMIT = 200
BULK_CELL_LIMIT = 20_000
CHUNK_ROW_SIZE = 100
ATTACHMENT_TABLE_WIDTH_PCT = "5000"
ATTACHMENT_TABLE_CONTENT_WIDTH_DXA = 12756
ATTACHMENT_COLUMN_MIN_WEIGHT = 2
LogFunc = Callable[..., None]
PAYMENT_TERMS_OVERRIDE_KEY = "paymentTermsOverride"
ITEMS_CONTENT_OVERRIDE_KEY = "itemsContentOverride"
ATTACHMENT_DETAIL_REF = "详情见附件"
TITLE_SCALAR_KEYS = ("purchaseSubject", "workDescription", "projectName", "engineeringScope")
TITLE_COLUMN_KEYS = ("name", "laborItem", "node")
DETAIL_COLUMN_KEYS = ("spec", "remark", "progressDescription")
AMOUNT_SCALAR_KEY = "totalAmount"
AMOUNT_COLUMN_KEY = "totalPrice"
NO_UNDERLINE_SCALAR_KEYS = {
    "signYear",
    "signMonth",
    "signDay",
    "signatureYear",
    "signatureMonth",
    "signatureDay",
    "supplierName",
    "supplierAddress",
    "supplierBank",
    "supplierAccount",
    "supplierTaxNo",
    "supplierPhone",
    "supplierFax",
    "supplierRepresentativeName",
    "supplierRepresentativeIdNo",
    "supplierRepresentativePhone",
    "supplierRepresentativeAddress",
    "supplierRepresentativeEmail",
    "buyerAuthorizedRepresentative",
    "supplierAuthorizedRepresentative",
    PAYMENT_TERMS_OVERRIDE_KEY,
    ITEMS_CONTENT_OVERRIDE_KEY,
}


def _stringify(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    return str(value)


def _pending(label: str) -> str:
    return f"【待填写：{label}】"


def _value_for_context(value: Any, label: str, blank_missing: bool = False) -> str:
    if value is None or str(value).strip() == "":
        if blank_missing:
            return ""
        return _pending(label)
    return str(value)


def _rich_text(
    value: Any,
    label: str,
    typography: TemplateTypography,
    blank_missing: bool = False,
    underline: bool = False,
) -> RichText:
    text = _value_for_context(value, label, blank_missing)
    if underline and blank_missing and not text.strip():
        text = UNDERLINE_BLANK
    rich_text = RichText()
    rich_text.add(
        text,
        font=typography.rich_text_font,
        size=typography.size_half_pt,
        underline=underline,
    )
    return rich_text


def _rich_text_multiline(
    value: Any,
    label: str,
    typography: TemplateTypography,
    blank_missing: bool = False,
) -> RichText:
    text = _value_for_context(value, label, blank_missing)
    rich_text = RichText()
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            rich_text.add("\n", font=typography.rich_text_font, size=typography.size_half_pt, underline=False)
        rich_text.add(line, font=typography.rich_text_font, size=typography.size_half_pt, underline=False)
    return rich_text


def _today_parts() -> dict[str, str]:
    now = datetime.now(ZoneInfo("Asia/Shanghai"))
    return {"year": f"{now.year}", "month": f"{now.month:02d}", "day": f"{now.day:02d}"}


def merge_render_data(patch: dict[str, Any], config: TemplateConfig) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for key in config.scalar_keys:
        out[key] = _stringify(patch.get(key))
    for table_name, columns in config.table_bindings.items():
        rows = patch.get(table_name)
        if not isinstance(rows, list):
            out[table_name] = []
            continue
        mapped_rows: list[dict[str, str]] = []
        for index, row in enumerate(rows):
            source = row if isinstance(row, dict) else {}
            mapped: dict[str, str] = {}
            for column in columns:
                fallback = str(index + 1) if column == "index" else ""
                mapped[column] = _stringify(source.get(column), fallback)
            mapped_rows.append(mapped)
        out[table_name] = mapped_rows
    parts = _today_parts()
    out.setdefault("signYear", parts["year"])
    out.setdefault("signMonth", parts["month"])
    out.setdefault("signDay", parts["day"])
    out.setdefault("signatureYear", parts["year"])
    out.setdefault("signatureMonth", parts["month"])
    out.setdefault("signatureDay", parts["day"])
    return out


def _first_non_empty_scalar(scalars: dict[str, Any], keys: tuple[str, ...]) -> str:
    for key in keys:
        value = _stringify(scalars.get(key)).strip()
        if value:
            return value
    return ""


def _first_column(columns: list[str], candidates: tuple[str, ...]) -> str | None:
    for candidate in candidates:
        if candidate in columns:
            return candidate
    return None


def build_attachment_summary_row(columns: list[str], scalars: dict[str, Any]) -> dict[str, str]:
    row = {column: "" for column in columns}
    if "index" in row:
        row["index"] = "1"
    title_column = _first_column(columns, TITLE_COLUMN_KEYS)
    if title_column:
        row[title_column] = _first_non_empty_scalar(scalars, TITLE_SCALAR_KEYS)
    detail_column = _first_column(columns, DETAIL_COLUMN_KEYS)
    if detail_column:
        row[detail_column] = ATTACHMENT_DETAIL_REF
    if AMOUNT_COLUMN_KEY in row:
        row[AMOUNT_COLUMN_KEY] = _stringify(scalars.get(AMOUNT_SCALAR_KEY)).strip()
    return row


def apply_attachment_table_summary(render_data: dict[str, Any], config: TemplateConfig) -> None:
    for table_name, columns in config.table_bindings.items():
        if not columns:
            continue
        render_data[table_name] = [build_attachment_summary_row(columns, render_data)]


def build_docxtpl_context(render_data: dict[str, Any], config: TemplateConfig, blank_missing: bool = False) -> dict[str, Any]:
    typography = get_template_typography(config.type)
    scalar_labels = {field["key"]: field["label"] for field in config.schema.get("scalars", [])}
    context: dict[str, Any] = {}
    override_text = _stringify(render_data.get(PAYMENT_TERMS_OVERRIDE_KEY)).strip()
    items_override_text = _stringify(render_data.get(ITEMS_CONTENT_OVERRIDE_KEY)).strip()
    context["hasPaymentTermsOverride"] = bool(override_text)
    context["hasItemsContentOverride"] = bool(items_override_text)
    for key in config.scalar_keys:
        if key in (PAYMENT_TERMS_OVERRIDE_KEY, ITEMS_CONTENT_OVERRIDE_KEY):
            continue
        context[key] = _rich_text(
            render_data.get(key),
            scalar_labels.get(key, key),
            typography,
            blank_missing,
            underline=key not in NO_UNDERLINE_SCALAR_KEYS,
        )
    context[PAYMENT_TERMS_OVERRIDE_KEY] = (
        _rich_text_multiline(
            override_text,
            scalar_labels.get(PAYMENT_TERMS_OVERRIDE_KEY, "付款期限覆盖内容"),
            typography,
            blank_missing,
        )
        if override_text
        else RichText()
    )
    context[ITEMS_CONTENT_OVERRIDE_KEY] = (
        _rich_text_multiline(
            items_override_text,
            scalar_labels.get(ITEMS_CONTENT_OVERRIDE_KEY, "协议内容覆盖"),
            typography,
            blank_missing,
        )
        if items_override_text
        else RichText()
    )
    for table_name, columns in config.table_bindings.items():
        table_def = config.schema.get("tables", {}).get(table_name, {})
        labels = {column["key"]: column["label"] for column in table_def.get("columns", [])}
        rows = render_data.get(table_name)
        mapped_rows: list[dict[str, str]] = []
        if isinstance(rows, list):
            for row in rows:
                source = row if isinstance(row, dict) else {}
                mapped_rows.append({
                    column: _rich_text(
                        source.get(column),
                        labels.get(column, column),
                        typography,
                        blank_missing,
                        underline=False,
                    )
                    for column in columns
                })
        context[table_name] = mapped_rows
    return context


def _non_empty_rows(rows: Any) -> list[list[str]]:
    if not isinstance(rows, list):
        return []
    normalized: list[list[str]] = []
    for row in rows:
        if not isinstance(row, list):
            continue
        cells = [_stringify(cell) for cell in row]
        if any(cell.strip() for cell in cells):
            normalized.append(cells)
    return normalized


def _elapsed_ms(start: float) -> int:
    return int((time.perf_counter() - start) * 1000)


def _log(logger: LogFunc | None, message: str, **meta: Any) -> None:
    if logger:
        logger(message, **meta)


def _set_run_font(run: Any, typography: TemplateTypography, size_pt: float | None = None) -> None:
    font_name = typography.east_asia
    run.font.name = font_name
    run.font.size = Pt(size_pt if size_pt is not None else typography.size_pt)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)


def _format_template_body_paragraphs(path: Path, template_type: str) -> None:
    typography = get_template_typography(template_type)
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            normalize_run_if_body(run, typography.east_asia, typography.size_half_pt)
    doc.save(str(path))


def _format_template_tables(path: Path, template_type: str) -> None:
    typography = get_template_typography(template_type)
    doc = Document(str(path))
    for table_index, table in enumerate(doc.tables):
        alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if table_index < LEFT_ALIGNED_HEADER_TABLE_COUNT
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignment
                    for run in paragraph.runs:
                        _set_run_font(run, typography)
    doc.save(str(path))


def _add_attachment_heading(doc: Any, text: str, level: int, typography: TemplateTypography) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    _set_run_font(run, typography, size_pt=14 if level == 1 else typography.size_pt)


def _set_table_grid_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _set_table_autofit_layout(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")


def choose_attachment_write_strategy(row_count: int, col_count: int) -> str:
    cell_count = row_count * col_count
    if cell_count <= SMALL_CELL_LIMIT:
        return "cell_api"
    if cell_count <= BULK_CELL_LIMIT:
        return "xml_bulk"
    return "xml_chunked"


def _append_table_borders(tbl_pr: OxmlElement) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def _append_table_autofit_layout(tbl_pr: OxmlElement) -> None:
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)


def _append_table_width(tbl_pr: OxmlElement) -> None:
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), ATTACHMENT_TABLE_WIDTH_PCT)
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def _append_table_layout(tbl_pr: OxmlElement, layout_type: str) -> None:
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), layout_type)
    tbl_pr.append(layout)


def _text_display_width(text: str) -> int:
    width = 0
    for char in str(text).strip():
        if unicodedata.east_asian_width(char) in {"F", "W"}:
            width += 2
        else:
            width += 1
    return width


def _column_display_weights(rows: list[list[str]], max_cols: int, min_weight: int = ATTACHMENT_COLUMN_MIN_WEIGHT) -> list[int]:
    header = rows[0] if rows else []
    weights: list[int] = []
    for col_index in range(max_cols):
        header_width = _text_display_width(header[col_index]) if col_index < len(header) else 0
        body_max = max(
            (_text_display_width(str(row[col_index])) for row in rows[1:] if col_index < len(row)),
            default=0,
        )
        weights.append(max(header_width, body_max, min_weight))
    return weights


def compute_attachment_column_widths(rows: list[list[str]], max_cols: int) -> list[int]:
    weights = _column_display_weights(rows, max_cols)
    return _proportional_column_widths(weights, ATTACHMENT_TABLE_CONTENT_WIDTH_DXA)


def _proportional_column_widths(weights: list[int], total_width_dxa: int) -> list[int]:
    if not weights:
        return []
    weight_total = sum(weights)
    if weight_total <= 0:
        equal = total_width_dxa // len(weights)
        return [equal] * len(weights)
    widths = [int(total_width_dxa * weight / weight_total) for weight in weights]
    remainder = total_width_dxa - sum(widths)
    if remainder:
        widest_index = max(range(len(widths)), key=widths.__getitem__)
        widths[widest_index] += remainder
    return widths


def _build_table_cell_xml(text: str) -> OxmlElement:
    tc = OxmlElement("w:tc")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    if text and (text[0].isspace() or text[-1].isspace()):
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    tc.append(paragraph)
    return tc


def _build_table_row_xml(row: list[str], max_cols: int) -> OxmlElement:
    tr = OxmlElement("w:tr")
    for col_index in range(max_cols):
        value = row[col_index] if col_index < len(row) else ""
        tr.append(_build_table_cell_xml(value))
    return tr


def _build_attachment_table_properties_xml() -> OxmlElement:
    tbl_pr = OxmlElement("w:tblPr")
    _append_table_width(tbl_pr)
    _append_table_layout(tbl_pr, "fixed")
    _append_table_borders(tbl_pr)
    return tbl_pr


def _build_table_grid_xml(rows: list[list[str]], max_cols: int) -> OxmlElement:
    col_widths = compute_attachment_column_widths(rows, max_cols)
    tbl_grid = OxmlElement("w:tblGrid")
    for width_dxa in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)
    return tbl_grid


def _build_table_xml(rows: list[list[str]], max_cols: int) -> OxmlElement:
    tbl = OxmlElement("w:tbl")
    tbl.append(_build_attachment_table_properties_xml())
    tbl.append(_build_table_grid_xml(rows, max_cols))
    for row in rows:
        tbl.append(_build_table_row_xml(row, max_cols))
    return tbl


def _append_table_rows_chunked(tbl: OxmlElement, rows: list[list[str]], max_cols: int, chunk_size: int) -> None:
    for start in range(0, len(rows), chunk_size):
        chunk = rows[start : start + chunk_size]
        for row in chunk:
            tbl.append(_build_table_row_xml(row, max_cols))


def _fill_table_with_cell_api(table: Any, rows: list[list[str]], max_cols: int) -> None:
    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def _apply_typography_to_table(table: Any, typography: TemplateTypography) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, typography)


def _append_attachment_table(
    doc: Any,
    rows: list[list[str]],
    strategy: str,
    typography: TemplateTypography,
) -> dict[str, int]:
    max_cols = max(len(row) for row in rows)
    row_count = len(rows)
    cell_count = row_count * max_cols
    if strategy == "cell_api":
        table = doc.add_table(rows=row_count, cols=max_cols)
        _set_table_autofit_layout(table)
        _set_table_grid_borders(table)
        _fill_table_with_cell_api(table, rows, max_cols)
        _apply_typography_to_table(table, typography)
    elif strategy == "xml_bulk":
        doc.element.body.append(_build_table_xml(rows, max_cols))
        _apply_typography_to_table(doc.tables[-1], typography)
    else:
        tbl = OxmlElement("w:tbl")
        tbl.append(_build_attachment_table_properties_xml())
        tbl.append(_build_table_grid_xml(rows, max_cols))
        _append_table_rows_chunked(tbl, rows, max_cols, CHUNK_ROW_SIZE)
        doc.element.body.append(tbl)
        _apply_typography_to_table(doc.tables[-1], typography)
    return {"rowCount": row_count, "colCount": max_cols, "cellCount": cell_count}


def _landscape_sect_pr(sect_pr: OxmlElement) -> OxmlElement:
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is not None:
        width = pg_sz.get(qn("w:w"))
        height = pg_sz.get(qn("w:h"))
        if width and height:
            pg_sz.set(qn("w:w"), height)
            pg_sz.set(qn("w:h"), width)
        pg_sz.set(qn("w:orient"), "landscape")
    return sect_pr


def _append_landscape_section_break(doc: Any) -> None:
    body = doc.element.body
    if body.sectPr is None:
        return
    landscape_sect = _landscape_sect_pr(deepcopy(body.sectPr))
    paragraph = doc.add_paragraph()
    paragraph_pr = OxmlElement("w:pPr")
    paragraph_pr.append(deepcopy(landscape_sect))
    paragraph._element.insert(0, paragraph_pr)
    body.replace(body.sectPr, landscape_sect)


def _strip_caigouhetong_attachment_section(path: Path, config: TemplateConfig, table_mode: str) -> None:
    if config.type != "caigouhetong" or table_mode == "attachment":
        return
    doc = Document(str(path))
    attachment_paragraph = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "附件一：设备采购清单"), None)
    if attachment_paragraph is None:
        return

    body = doc.element.body
    attachment_element = attachment_paragraph._element
    previous_element = attachment_element.getprevious()
    previous_section = None
    if previous_element is not None and previous_element.tag == qn("w:p"):
        previous_props = previous_element.find(qn("w:pPr"))
        if previous_props is not None:
            previous_section = previous_props.find(qn("w:sectPr"))
    if previous_section is not None and body.sectPr is not None:
        body.replace(body.sectPr, deepcopy(previous_section))
        previous_section.getparent().remove(previous_section)

    removable_elements = [attachment_element]
    cursor = attachment_element.getnext()
    while cursor is not None and cursor.tag == qn("w:p"):
        texts = [node.text or "" for node in cursor.iter(qn("w:t"))]
        if any(text.strip() for text in texts):
            break
        removable_elements.append(cursor)
        cursor = cursor.getnext()
    for element in removable_elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    doc.save(str(path))


def append_quote_attachment(
    path: Path,
    quote_attachment: dict[str, Any] | None,
    typography: TemplateTypography,
    logger: LogFunc | None = None,
) -> None:
    sheets = quote_attachment.get("sheets") if isinstance(quote_attachment, dict) else None
    if not isinstance(sheets, list):
        return
    valid_sheets: list[tuple[str, list[list[str]]]] = []
    for sheet in sheets:
        rows = _non_empty_rows(sheet.get("rows") if isinstance(sheet, dict) else None)
        if not rows:
            continue
        sheet_name = str(sheet.get("name") or "Sheet") if isinstance(sheet, dict) else "Sheet"
        valid_sheets.append((sheet_name, rows))
    if not valid_sheets:
        return

    start = time.perf_counter()
    doc = Document(str(path))
    doc.add_page_break()
    _append_landscape_section_break(doc)
    _add_attachment_heading(doc, "附件：报价单明细", level=1, typography=typography)
    sheet_stats: list[dict[str, Any]] = []
    for sheet_name, rows in valid_sheets:
        _add_attachment_heading(doc, sheet_name, level=2, typography=typography)
        max_cols = max(len(row) for row in rows)
        strategy = choose_attachment_write_strategy(len(rows), max_cols)
        table_stats = _append_attachment_table(doc, rows, strategy, typography)
        sheet_stats.append({"sheetName": sheet_name, "writeStrategy": strategy, **table_stats})
    doc.save(str(path))
    _log(
        logger,
        "quote attachment appended",
        outputFile=path.name,
        sheetCount=len(valid_sheets),
        sheetStats=sheet_stats,
        elapsedMs=_elapsed_ms(start),
    )


def render_contract(
    render_data: dict[str, Any],
    config: TemplateConfig,
    contract_id: str,
    blank_missing: bool = False,
    table_mode: str = "auto",
    quote_attachment: dict[str, Any] | None = None,
    logger: LogFunc | None = None,
) -> Path:
    ensure_storage()
    template_path = template_docx_path(config.type)
    typography = get_template_typography(config.type)
    output_path = CONTRACTS_DIR / f"{contract_id}.docx"
    if table_mode == "attachment" and config.table_bindings:
        apply_attachment_table_summary(render_data, config)
    template_start = time.perf_counter()
    doc = DocxTemplate(str(template_path))
    doc.render(build_docxtpl_context(render_data, config, blank_missing=blank_missing))
    doc.save(str(output_path))
    _strip_caigouhetong_attachment_section(output_path, config, table_mode)
    _format_template_body_paragraphs(output_path, config.type)
    _format_template_tables(output_path, config.type)
    _log(logger, "contract template rendered", outputFile=output_path.name, templateType=config.type, elapsedMs=_elapsed_ms(template_start))
    append_quote_attachment(output_path, quote_attachment, typography, logger)
    return output_path
