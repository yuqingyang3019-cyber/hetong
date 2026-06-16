"""Normalize body typography in contract template docx files."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from agent.contract.config import TEMPLATE_BASENAME, TEMPLATE_ROOT, get_template_typography
from agent.scripts.template_docx_utils import normalize_run_if_body

BASENAME_TO_TYPE = {basename: template_type for template_type, basename in TEMPLATE_BASENAME.items()}


def normalize_document(path: Path, template_type: str) -> None:
    typography = get_template_typography(template_type)
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            normalize_run_if_body(run, typography.east_asia, typography.size_half_pt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        normalize_run_if_body(run, typography.east_asia, typography.size_half_pt)
    doc.save(str(path))


def main() -> None:
    for docx_path in sorted(TEMPLATE_ROOT.glob("*.docx")):
        if docx_path.name.endswith(".source.docx"):
            continue
        template_type = BASENAME_TO_TYPE.get(docx_path.stem)
        if template_type is None:
            continue
        normalize_document(docx_path, template_type)
        print(f"Normalized {docx_path.name} ({template_type})")


if __name__ == "__main__":
    main()
