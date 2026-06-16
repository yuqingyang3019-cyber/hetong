"""Patch caigouhetong payment override block and paragraph indentation."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from agent.scripts.template_docx_utils import copy_run_format, first_formatted_run

DOCX_PATH = ROOT / "agent" / "contract" / "templates" / "zhanweifu" / "caigouhetong.docx"


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    format_source: Paragraph | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_run = new_para.add_run(text)
        source_run = first_formatted_run(format_source) if format_source is not None else None
        if source_run is not None:
            copy_run_format(source_run, new_run)
    return new_para


def find_paragraph(doc: Document, exact_text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    return None


def ensure_override_block(doc: Document) -> None:
    if find_paragraph(doc, "{% if hasPaymentTermsOverride %}") and find_paragraph(doc, "{{r paymentTermsOverride }}"):
        return

    heading_idx = start_idx = end_idx = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "付款期限" and heading_idx is None:
            heading_idx = index
        if text.startswith("（1）预付款"):
            start_idx = index
        if text.startswith("（5）质保金"):
            end_idx = index
    if heading_idx is None or start_idx is None or end_idx is None:
        raise RuntimeError("未找到付款期限段落，模板结构可能已变更")

    heading_p = doc.paragraphs[heading_idx]
    end_p = doc.paragraphs[end_idx]
    format_source = doc.paragraphs[start_idx]

    if_block = insert_paragraph_after(heading_p, "{% if hasPaymentTermsOverride %}", format_source)
    override_p = insert_paragraph_after(if_block, "{{r paymentTermsOverride }}", format_source)
    insert_paragraph_after(override_p, "{% else %}", format_source)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("（5）质保金"):
            end_p = paragraph
            break
    insert_paragraph_after(end_p, "{% endif %}", format_source)


def sync_override_indent(doc: Document) -> None:
    override_p = find_paragraph(doc, "{{r paymentTermsOverride }}")
    first_payment_item = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("（1）预付款"):
            first_payment_item = paragraph
            break
    if override_p is None or first_payment_item is None:
        raise RuntimeError("未找到付款期限覆盖段或默认付款条款段")
    override_format = override_p.paragraph_format
    source_format = first_payment_item.paragraph_format
    override_format.left_indent = source_format.left_indent
    override_format.first_line_indent = source_format.first_line_indent


def main() -> None:
    doc = Document(str(DOCX_PATH))
    ensure_override_block(doc)
    sync_override_indent(doc)
    doc.save(str(DOCX_PATH))
    print(f"Patched {DOCX_PATH}")


if __name__ == "__main__":
    main()
