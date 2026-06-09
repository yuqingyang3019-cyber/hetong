"""Patch simple-contract.source.docx with docxtpl placeholders, preserving WPS layout."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Row
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_DIR = ROOT / "agent" / "contract" / "templates" / "zhanweifu"
SOURCE_PATH = TEMPLATE_DIR / "simple-contract.source.docx"
OUTPUT_PATH = TEMPLATE_DIR / "simple-contract.docx"
ROOT_SOURCE = ROOT / "简易合同模板.docx"
PLACEHOLDER_MARKER = "{{r contractNo }}"


def ensure_source() -> None:
    if not SOURCE_PATH.exists():
        if not ROOT_SOURCE.exists():
            raise FileNotFoundError(f"缺少母版文件：{SOURCE_PATH} 或 {ROOT_SOURCE}")
        shutil.copy2(ROOT_SOURCE, SOURCE_PATH)


def is_already_patched(doc: Document) -> bool:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if PLACEHOLDER_MARKER in cell.text:
                    return True
    for paragraph in doc.paragraphs:
        if PLACEHOLDER_MARKER in paragraph.text:
            return True
    return False


def append_to_paragraph(paragraph: Paragraph, suffix: str) -> None:
    if paragraph.runs:
        paragraph.runs[-1].text = f"{paragraph.runs[-1].text}{suffix}"
        return
    paragraph.add_run(suffix)


def find_paragraph(cell, predicate) -> Paragraph:
    for paragraph in cell.paragraphs:
        if predicate(paragraph.text):
            return paragraph
    raise RuntimeError(f"未找到匹配段落：{predicate}")


def replace_paragraph_text(paragraph: Paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        raise RuntimeError(f"段落中未找到锚点 {old!r}：{paragraph.text!r}")
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = paragraph.runs[0].text.replace(old, new, 1)
        return
    combined = "".join(run.text for run in paragraph.runs)
    if old not in combined:
        raise RuntimeError(f"段落 run 中未找到锚点 {old!r}")
    updated = combined.replace(old, new, 1)
    for index, run in enumerate(paragraph.runs):
        run.text = updated if index == 0 else ""


def replace_in_cell_lines(cell, replacements: list[tuple[str, str]]) -> None:
    for old, new in replacements:
        matched = False
        for paragraph in cell.paragraphs:
            if old in paragraph.text:
                replace_paragraph_text(paragraph, old, new)
                matched = True
                break
        if not matched:
            raise RuntimeError(f"单元格中未找到锚点 {old!r}：{cell.text!r}")


def clone_row_after(table, row_index: int) -> _Row:
    source_row = table.rows[row_index]
    new_tr = deepcopy(source_row._tr)
    source_row._tr.addnext(new_tr)
    return _Row(new_tr, table)


def patch_table0(doc: Document) -> None:
    table = doc.tables[0]
    replace_in_cell_lines(table.rows[0].cells[0], [("合同编号:", "合同编号:{{r contractNo }}")])
    replace_in_cell_lines(
        table.rows[1].cells[0],
        [("签订日期: 年 月 日", "签订日期: {{r signYear }} 年 {{r signMonth }} 月 {{r signDay }} 日")],
    )


def patch_table1(doc: Document) -> None:
    left = doc.tables[1].rows[0].cells[0]
    right = doc.tables[1].rows[0].cells[1]
    replace_in_cell_lines(
        left,
        [
            ("电话（Tel）： ", "电话（Tel）：{{r buyerPhone }} "),
            ("传真（Fax）： ", "传真（Fax）：{{r buyerFax }} "),
            ("联系人：", "联系人：{{r buyerContact }}"),
        ],
    )
    replace_in_cell_lines(
        right,
        [
            ("供方（乙方）：", "供方（乙方）：{{r supplierName }}"),
            ("地址：", "地址：{{r supplierAddress }}"),
            ("开户银行：", "开户银行：{{r supplierBank }}"),
            ("银行帐号：", "银行帐号：{{r supplierAccount }}"),
            ("行号：", "行号：{{r supplierBankCode }}"),
            ("税号：", "税号：{{r supplierTaxNo }}"),
            ("电话（Tel）： ", "电话（Tel）：{{r supplierPhone }} "),
            ("传真（Fax）：", "传真（Fax）：{{r supplierFax }}"),
            ("联系人：", "联系人：{{r supplierContact }}"),
        ],
    )


def patch_table2_items(doc: Document) -> None:
    table = doc.tables[2]
    if table.rows[0].cells[1].text.strip() != "产品名称":
        raise RuntimeError("TABLE2 表头结构已变更")

    table.rows[1].cells[0].text = "{%tr for item in items %}"
    data_row = clone_row_after(table, 1)
    item_columns = [
        "{{r item.index }}",
        "{{r item.name }}",
        "{{r item.spec }}",
        "{{r item.quantity }}",
        "{{r item.unit }}",
        "{{r item.unitPrice }}",
        "{{r item.totalPrice }}",
        "{{r item.remark }}",
    ]
    for index, placeholder in enumerate(item_columns):
        data_row.cells[index].text = placeholder

    end_row = clone_row_after(table, 2)
    end_row.cells[0].text = "{%tr endfor %}"


def patch_table2_amounts(doc: Document) -> None:
    cell = doc.tables[2].rows[4].cells[2]
    replace_in_cell_lines(
        cell,
        [
            (
                "（大写）人民币元整（￥.00元）",
                "（大写）人民币{{r totalAmountChinese }}元整（￥{{r totalAmount }}元）",
            ),
            ("不计税总金额：元", "不计税总金额：{{r amountWithoutTax }}元"),
        ],
    )


def patch_table2_delivery(doc: Document) -> None:
    cell = doc.tables[2].rows[5].cells[0]
    replace_in_cell_lines(
        cell,
        [
            ("含13%增值税", "含{{r taxRate }}%增值税"),
            ("交货期为： ", "交货期为：{{r deliveryDays }} "),
        ],
    )


def patch_paragraphs(doc: Document) -> None:
    delivery_patched = settlement_patched = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "交（提）货地点、方式及货物接收人：" in text and not delivery_patched:
            replace_paragraph_text(
                paragraph,
                "交（提）货地点、方式及货物接收人： 。",
                "交（提）货地点、方式及货物接收人：{{r deliveryPlace }}、{{r deliveryMethod }}及{{r goodsRecipient }}。",
            )
            delivery_patched = True
        if text.startswith("九、结算方式及期限：") and not settlement_patched:
            replace_paragraph_text(
                paragraph,
                "九、结算方式及期限：  。",
                "九、结算方式及期限：{{r settlementTerms }}。",
            )
            settlement_patched = True
    if not delivery_patched:
        raise RuntimeError("未找到交货地点段落")
    if not settlement_patched:
        raise RuntimeError("未找到结算方式段落")


def patch_table3(doc: Document) -> None:
    cell = doc.tables[3].rows[0].cells[1]
    for paragraph in cell.paragraphs:
        if "单位名称（章）：" in paragraph.text:
            replace_paragraph_text(paragraph, "单位名称（章）： ", "单位名称（章）：{{r supplierName }} ")
            return
    raise RuntimeError("未找到供方落款段落")


def patch_document(doc: Document) -> None:
    if is_already_patched(doc):
        return
    patch_table0(doc)
    patch_table1(doc)
    patch_table2_items(doc)
    patch_table2_amounts(doc)
    patch_table2_delivery(doc)
    patch_paragraphs(doc)
    patch_table3(doc)


def main() -> None:
    ensure_source()
    doc = Document(str(SOURCE_PATH))
    patch_document(doc)
    doc.save(str(OUTPUT_PATH))
    print(f"Patched {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
