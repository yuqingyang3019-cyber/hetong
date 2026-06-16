"""Shared helpers for patching and normalizing contract template docx files."""
from __future__ import annotations

from copy import deepcopy
from typing import Any
from xml.etree import ElementTree as ET

from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NS = {"w": W_NS}
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

HEADING_SIZE_HALF_PT = 28


def run_size_half_pt(run: Any) -> int | None:
    r_pr = run._element.rPr
    if r_pr is None:
        return None
    sz_el = r_pr.find(qn("w:sz"))
    if sz_el is None:
        return None
    value = sz_el.get(qn("w:val"))
    return int(value) if value else None


def is_heading_run(run: Any) -> bool:
    size = run_size_half_pt(run)
    return size is not None and size >= HEADING_SIZE_HALF_PT


def apply_run_typography(run: Any, east_asia: str, size_half_pt: int) -> None:
    run.font.name = east_asia
    run.font.size = Pt(size_half_pt / 2)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), east_asia)
    r_fonts.set(qn("w:hAnsi"), east_asia)
    sz = r_pr.get_or_add_sz()
    sz.set(qn("w:val"), str(size_half_pt))


def normalize_run_if_body(run: Any, east_asia: str, size_half_pt: int) -> None:
    if is_heading_run(run):
        return
    apply_run_typography(run, east_asia, size_half_pt)


def copy_run_format(source_run: Any, target_run: Any) -> None:
    source_r_pr = source_run._element.rPr
    if source_r_pr is None:
        return
    target_element = target_run._element
    existing = target_element.rPr
    if existing is not None:
        target_element.remove(existing)
    target_element.insert(0, deepcopy(source_r_pr))


def replace_paragraph_text_preserving_format(paragraph: Paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        raise RuntimeError(f"段落中未找到锚点 {old!r}：{paragraph.text!r}")
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = paragraph.runs[0].text.replace(old, new, 1)
        return
    combined = "".join(run.text for run in paragraph.runs)
    if old not in combined:
        raise RuntimeError(f"段落 run 中未找到锚点 {old!r}")
    updated = combined.replace(old, new, 1)
    for index, run in enumerate(paragraph.runs):
        run.text = updated if index == 0 else ""


def append_to_paragraph_preserving_format(paragraph: Paragraph, suffix: str) -> None:
    if paragraph.runs:
        paragraph.runs[-1].text = f"{paragraph.runs[-1].text}{suffix}"
        return
    paragraph.add_run(suffix)


def set_paragraph_text_preserving_format(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for extra_run in paragraph.runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)
        first_run.text = text
        return
    paragraph.add_run(text)


def set_cell_text_preserving_format(cell: Any, text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text_preserving_format(cell.paragraphs[0], text)
        return
    cell.text = text


def first_formatted_run(paragraph: Paragraph) -> Any | None:
    for run in paragraph.runs:
        if run.text.strip() or run._element.rPr is not None:
            return run
    return paragraph.runs[0] if paragraph.runs else None


def set_cell_text_preserving_format_xml(cell: ET.Element, text: str) -> None:
    paragraphs = cell.findall("w:p", NS)
    paragraph = paragraphs[0] if paragraphs else ET.SubElement(cell, f"{W}p")
    runs = paragraph.findall("w:r", NS)
    preserved_r_pr = None
    if runs:
        source_r_pr = runs[0].find("w:rPr", NS)
        if source_r_pr is not None:
            preserved_r_pr = deepcopy(source_r_pr)
    for run in list(paragraph.findall("w:r", NS)):
        paragraph.remove(run)
    run = ET.SubElement(paragraph, f"{W}r")
    if preserved_r_pr is not None:
        run.insert(0, preserved_r_pr)
    text_node = ET.SubElement(run, f"{W}t")
    text_node.text = text
    if text and (text[0] == " " or text[-1] == " "):
        text_node.set(XML_SPACE, "preserve")


def make_paragraph_xml(text: str, format_from_run: ET.Element | None = None) -> ET.Element:
    paragraph = ET.Element(f"{W}p")
    run = ET.SubElement(paragraph, f"{W}r")
    if format_from_run is not None:
        source_r_pr = format_from_run.find("w:rPr", NS)
        if source_r_pr is not None:
            run.insert(0, deepcopy(source_r_pr))
    text_node = ET.SubElement(run, f"{W}t")
    text_node.text = text
    return paragraph


def first_run_in_element(element: ET.Element) -> ET.Element | None:
    for run in element.iter(f"{W}r"):
        return run
    return None
