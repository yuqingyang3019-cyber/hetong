"""Build simple-contract.docx with docxtpl placeholders from the source .doc content."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "agent" / "contract" / "templates" / "zhanweifu" / "simple-contract.docx"


def add_paragraph(doc: Document, text: str, *, bold: bool = False, align_center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def add_table_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for index, text in enumerate(cells):
        row.cells[index].text = text


def main() -> None:
    doc = Document()

    add_paragraph(doc, "产品购销合同", bold=True, align_center=True)
    add_paragraph(doc, "合同编号:{{r contractNo }}")
    add_paragraph(doc, "签订地点：杭州市余杭区")
    add_paragraph(doc, "签订日期: {{r signYear }} 年 {{r signMonth }} 月 {{r signDay }} 日")
    doc.add_paragraph()

    add_paragraph(doc, "需方（甲方）：浙江沃乐科技股份有限公司")
    add_paragraph(doc, "地址： 浙江省杭州市余杭区仓前街道仓兴街397号18幢5层")
    add_paragraph(doc, "开户银行：上海浦东发展银行股份有限公司余杭科技城支行")
    add_paragraph(doc, "银行帐号：95260078801100000902")
    add_paragraph(doc, "税号： 91330110MA2HYAEK1T(1/1)")
    add_paragraph(doc, "电话（Tel）：{{r buyerPhone }}")
    add_paragraph(doc, "传真（Fax）：{{r buyerFax }}")
    add_paragraph(doc, "联系人：{{r buyerContact }}")

    add_paragraph(doc, "供方（乙方）：{{r supplierName }}")
    add_paragraph(doc, "地址：{{r supplierAddress }}")
    add_paragraph(doc, "开户银行：{{r supplierBank }}")
    add_paragraph(doc, "银行帐号：{{r supplierAccount }}")
    add_paragraph(doc, "行号：{{r supplierBankCode }}")
    add_paragraph(doc, "税号：{{r supplierTaxNo }}")
    add_paragraph(doc, "电话（Tel）：{{r supplierPhone }}")
    add_paragraph(doc, "传真（Fax）：{{r supplierFax }}")
    add_paragraph(doc, "联系人：{{r supplierContact }}")

    add_paragraph(doc, "产品名称、型号、数量、金额及交货时间等；")

    table = doc.add_table(rows=4, cols=8)
    table.style = "Table Grid"
    headers = ["序号", "产品名称", "规格", "数量", "单位", "单价（元）", "小计（元）", "备注"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    table.rows[1].cells[0].text = "{%tr for item in items %}"
    data_cells = [
        "{{r item.index }}",
        "{{r item.name }}",
        "{{r item.spec }}",
        "{{r item.quantity }}",
        "{{r item.unit }}",
        "{{r item.unitPrice }}",
        "{{r item.totalPrice }}",
        "{{r item.remark }}",
    ]
    for index, text in enumerate(data_cells):
        table.rows[2].cells[index].text = text
    table.rows[3].cells[0].text = "{%tr endfor %}"

    add_paragraph(doc, "最终优惠价")
    add_paragraph(doc, "（大写）人民币{{r totalAmountChinese }}元整（￥{{r totalAmount }}元）")
    add_paragraph(doc, "不计税总金额：{{r amountWithoutTax }}元")
    doc.add_paragraph()
    add_paragraph(
        doc,
        "备注：此价格含{{r taxRate }}%增值税、运费。                         交货期为：{{r deliveryDays }}",
    )
    add_paragraph(
        doc,
        "二、质量要求技术标准，供方对质量负责的条件和期限：详细技术要求按国家有关标准及合同约定技术要求为准，质保期12个月。",
    )
    add_paragraph(
        doc,
        "三、交（提）货地点、方式及货物接收人：{{r deliveryPlace }}、{{r deliveryMethod }}及{{r goodsRecipient }}。",
    )
    add_paragraph(doc, "四、运输方式及到达港和费用负担：货物送至指定地点，费用由供方提供，送货提前联系收件人。")
    add_paragraph(doc, "五、合理损耗及计算方法：/")
    add_paragraph(doc, "六、包装标准、方法及提出异议期限：按出厂标准包装；")
    add_paragraph(
        doc,
        "七、验收标准、方法及提出异议期限：交货时按本合同及合同附件清单进行规格、数量、外观检查和验收，验收有异议应3天内提出并要求更换合格的产品；",
    )
    add_paragraph(doc, "八、随机备品备件、工具数量及供应方式：无。")
    add_paragraph(doc, "九、结算方式及期限：{{r settlementTerms }}。项目过程中另有增补项，按本合同单价按实结算。")
    add_paragraph(doc, "十、如需提供担保，另立合同担保书，作为本合同附件：/")
    add_paragraph(doc, "十一、违约责任：违约项按照有关经济合同法规的条款，协商解决。")
    add_paragraph(doc, "十二、本合同在履行过程中发生争议，由当事人双方协商解决；若协商不成，可向合同签订地人民法院起诉。")
    add_paragraph(doc, "十三、合同附件，与本合同具有同等法律效力。")
    add_paragraph(doc, "其他约定事项：本合同一式贰份，经双方盖章后生效。")
    doc.add_paragraph()
    add_paragraph(doc, "需      方")
    add_paragraph(doc, "单位名称（章）：浙江沃乐科技股份有限公司")
    doc.add_paragraph()
    add_paragraph(doc, "供        方")
    add_paragraph(doc, "单位名称（章）：{{r supplierName }}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
