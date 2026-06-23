from __future__ import annotations

from pathlib import Path
import base64
import json
import os
import re
import sys
import types
from io import BytesIO
from datetime import date
from typing import Any
from unittest.mock import ANY, Mock, patch
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.testclient import TestClient

from agent.contract import llm as contract_llm
from agent.contract.config import (
    MAX_STORED_FILE_NAME_BYTES,
    TEMPLATE_BASENAME,
    UPLOADS_DIR,
    get_template_config,
    get_template_typography,
    template_docx_path,
    truncate_file_name_to_bytes,
)
from agent.contract.extract import extract_excel_payload, extract_excel_text, extract_pdf_text
from agent.contract.render import (
    ATTACHMENT_TABLE_CONTENT_WIDTH_DXA,
    append_quote_attachment,
    apply_attachment_table_summary,
    build_attachment_summary_row,
    build_docxtpl_context,
    choose_attachment_write_strategy,
    compute_attachment_column_widths,
    merge_render_data,
    render_contract,
    _text_display_width,
)
from agent.main import STATIC_DIR, app, apply_delivery_date_calculation, apply_line_item_calculations, apply_tax_calculations, contract_download_payload, contract_file_name, generate_contract, sign_session_payload
from agent.yonyou_vendor import (
    EXPLICIT_VENDOR_DATA_FIELDS,
    apply_yonbip_supplier_patch,
    supplier_patch_from_yonbip,
    vendor_lookup_payload,
    vendor_query_payload,
)


os.environ.setdefault("APP_SESSION_SECRET", "test-session-secret-123456789012")

client = TestClient(app)
PNG_BYTES = b"\x89PNG\r\n\x1a\nquote-image"
JPEG_BYTES = b"\xff\xd8\xff\xe0\x00\x10JFIF\x00quote-image"
BMP_BYTES = b"BMquote-image"
GIF_BYTES = b"GIF89aquote-image"
TIFF_BYTES = b"II*\x00quote-image"
WEBP_BYTES = b"RIFF\x0c\x00\x00\x00WEBPquote-image"


def xlsx_bytes(rows: list[list[object]], sheet_name: str = "报价") -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    for row in rows:
        sheet.append(row)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def agent_auth_header(userid: str = "uid1", unionid: str = "union-x") -> dict[str, str]:
    token = sign_session_payload({
        "typ": "agent",
        "exp": 4_102_444_800,
        "userid": userid,
        "name": "张三",
        "unionid": unionid,
    })
    return {"Authorization": f"Bearer {token}"}


def upload_quote(
    original_name: str = "quote.pdf",
    mime_type: str = "application/pdf",
    content: bytes = b"%PDF-1.4 quote",
    headers: dict[str, str] | None = None,
) -> dict:
    response = client.post(
        "/api/uploads",
        headers=headers or agent_auth_header(),
        json={
            "originalName": original_name,
            "mimeType": mime_type,
            "size": len(content),
            "data": base64.b64encode(content).decode("ascii"),
        },
    )
    assert response.status_code == 200
    return response.json()


def upload_record(upload_id: str) -> dict:
    return json.loads((UPLOADS_DIR / f"{upload_id}.json").read_text(encoding="utf-8"))


def test_health() -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"ok": True}


def test_h5_page_returns_not_found_without_static_build() -> None:
    response = client.get("/h5")
    assert response.status_code == (200 if STATIC_DIR.exists() else 404)


def test_upload_multipart_pdf() -> None:
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        files={"file": ("quote.pdf", b"%PDF-1.4 quote", "application/pdf")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["id"]
    assert body["ok"] is True
    assert "path" not in body


def test_upload_json_base64_pdf() -> None:
    content = b"%PDF-1.4 hello quote"
    body = upload_quote(content=content)
    assert body["size"] == len(content)
    assert body["ok"] is True


LONG_UPLOAD_FILENAME = (
    "浙江沃乐科技股份有限公司-禾润肯尼亚-电磁流量计询价单--杭州美仪报价2026.4.21-"
    "改：最终优惠价23000元（改：中水部分取消；FIT-501数量3改为1；FIT-202参数变更为："
    "Q=0-40m3h，DN65，FIT-202 调整为铂金电极） - 副本.xlsx"
)


def test_truncate_file_name_to_bytes_preserves_extension() -> None:
    long_name = f"{'浙' * 80}.xlsx"
    truncated = truncate_file_name_to_bytes(long_name, 100)
    assert truncated.endswith(".xlsx")
    assert len(truncated.encode("utf-8")) <= 100


def test_upload_long_chinese_filename_succeeds() -> None:
    content = xlsx_bytes([["品名", "数量"], ["设备1", 1]])
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": LONG_UPLOAD_FILENAME,
            "mimeType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "size": len(content),
            "data": base64.b64encode(content).decode("ascii"),
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["ok"] is True
    assert body["originalName"] == LONG_UPLOAD_FILENAME
    record = upload_record(body["id"])
    assert Path(record["path"]).exists()
    assert len(record["fileName"].encode("utf-8")) <= MAX_STORED_FILE_NAME_BYTES
    assert record["fileName"].endswith(".xlsx")


def test_upload_rejects_extremely_long_filename() -> None:
    original_name = f"{'浙' * 200}.xlsx"
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": original_name,
            "mimeType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "size": 4,
            "data": base64.b64encode(xlsx_bytes([["品名", "数量"], ["设备1", 1]])).decode("ascii"),
        },
    )
    assert response.status_code == 400
    body = response.json()
    assert body["code"] == "FILENAME_TOO_LONG"


@pytest.mark.parametrize(
    ("original_name", "mime_type", "content", "expected_mime_type"),
    [
        ("quote.jpg", "image/jpeg", JPEG_BYTES, "image/jpeg"),
        ("quote.png", "image/png", PNG_BYTES, "image/png"),
        ("quote.bmp", "image/bmp", BMP_BYTES, "image/bmp"),
        ("quote.gif", "image/gif", GIF_BYTES, "image/gif"),
        ("quote.tif", "image/tiff", TIFF_BYTES, "image/tiff"),
        ("quote.webp", "image/webp", WEBP_BYTES, "image/webp"),
    ],
)
def test_upload_accepts_common_image_formats(
    original_name: str,
    mime_type: str,
    content: bytes,
    expected_mime_type: str,
) -> None:
    body = upload_quote(original_name=original_name, mime_type=mime_type, content=content)

    assert body["ok"] is True
    assert body["mimeType"] == expected_mime_type
    assert upload_record(body["id"])["mimeType"] == expected_mime_type


def test_upload_accepts_mismatched_supported_image_extension() -> None:
    body = upload_quote(original_name="quote.jpg", mime_type="image/jpeg", content=PNG_BYTES)

    assert body["ok"] is True
    assert body["mimeType"] == "image/png"


def test_upload_rejects_txt() -> None:
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.txt",
            "mimeType": "text/plain",
            "size": len(b"hello quote"),
            "data": base64.b64encode(b"hello quote").decode("ascii"),
        },
    )
    assert response.status_code == 400
    body = response.json()
    assert body["code"] == "UNSUPPORTED_FILE_TYPE"


def test_upload_rejects_unknown_image_signature() -> None:
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.webp",
            "mimeType": "image/webp",
            "size": len(b"not-an-image"),
            "data": base64.b64encode(b"not-an-image").decode("ascii"),
        },
    )

    assert response.status_code == 400
    body = response.json()
    assert body["code"] == "INVALID_ARGUMENT"
    assert "文件内容与格式不匹配" in body["message"]


def test_upload_rejects_mismatched_file_signature() -> None:
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.xls",
            "mimeType": "application/vnd.ms-excel",
            "size": len(PNG_BYTES),
            "data": base64.b64encode(PNG_BYTES).decode("ascii"),
        },
    )

    assert response.status_code == 400
    body = response.json()
    assert body["code"] == "INVALID_ARGUMENT"
    assert "文件内容与格式不匹配" in body["message"]


def test_extract_excel_text_outputs_tsv_without_html(tmp_path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "报价"
    sheet.append(["品名", "数量", "单价"])
    sheet.append(["阀门", 2, 100])
    path = tmp_path / "quote.xlsx"
    workbook.save(path)

    text = extract_excel_text(path)

    assert "[表格 parser=excel sheet=报价 format=tsv]" in text
    assert "品名\t数量\t单价" in text
    assert "阀门\t2\t100" in text
    assert "<table" not in text
    assert "<td>" not in text


def test_extract_excel_payload_marks_complex_attachment_mode(tmp_path: Path) -> None:
    content = xlsx_bytes([["品名", "数量"], *[[f"设备{i}", i] for i in range(1, 6)]])
    path = tmp_path / "quote.xlsx"
    path.write_bytes(content)

    payload = extract_excel_payload(path)

    assert payload["attachmentMode"]["enabled"] is True
    assert payload["attachmentMode"]["rowCount"] == 6
    assert "total_rows_over_threshold" in payload["attachmentMode"]["reasons"]
    assert payload["sheets"][0]["name"] == "报价"


def test_extract_excel_payload_marks_multi_sheet_attachment_mode(tmp_path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.active.title = "报价一"
    workbook.active.append(["品名", "数量"])
    workbook.active.append(["阀门", 1])
    sheet = workbook.create_sheet("报价二")
    sheet.append(["品名", "数量"])
    sheet.append(["水泵", 2])
    path = tmp_path / "multi.xlsx"
    workbook.save(path)

    payload = extract_excel_payload(path)

    assert payload["attachmentMode"]["enabled"] is True
    assert payload["attachmentMode"]["sheetCount"] == 2
    assert "multiple_sheets" in payload["attachmentMode"]["reasons"]


def test_extract_pdf_text_outputs_tsv_without_html() -> None:
    class FakeTable:
        def extract(self) -> list[list[str]]:
            return [["品名", "数量"], ["阀门", "2"]]

    class FakePage:
        def find_tables(self, table_settings: dict) -> list[FakeTable]:
            return [FakeTable()]

        def extract_text(self, **kwargs: object) -> str:
            return "报价备注"

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, exc_type: object, exc: object, traceback: object) -> None:
            return None

    with patch("pdfplumber.open", return_value=FakePdf()):
        text = extract_pdf_text(Path("quote.pdf"))

    assert "[表格 parser=pdfplumber-lines page=1 index=1 format=tsv]" in text
    assert "品名\t数量" in text
    assert "阀门\t2" in text
    assert "[第 1 页文字]" in text
    assert "<table" not in text
    assert "<td>" not in text


def test_yonbip_vendor_query_uses_explicit_fields() -> None:
    payload = vendor_query_payload(1, 500)

    assert payload["data"] == EXPLICIT_VENDOR_DATA_FIELDS
    assert "vendorbanks" in payload["partParam"]
    assert "vendorcontactss" in payload["partParam"]


def test_yonbip_vendor_lookup_filters_by_supplier_name() -> None:
    payload = vendor_lookup_payload("供应商A", 1, 10)

    assert payload["data"] == EXPLICIT_VENDOR_DATA_FIELDS
    assert payload["condition"]["simpleVOs"] == [{"field": "name", "op": "eq", "value1": "供应商A"}]
    assert payload["partParam"]["vendorbanks"]["data"] == "*,openaccountbank.name"
    assert payload["partParam"]["vendorcontactss"]["data"] == "*"


def test_supplier_patch_from_yonbip_overwrites_title_fields() -> None:
    extracted = {
        "supplierName": "供应商A",
        "supplierAddress": "报价单地址",
        "supplierBank": "报价单开户行",
        "supplierAccount": None,
    }
    records = [{
        "code": "S001",
        "name": "用友供应商A",
        "creditcode": "9133",
        "address": "用友地址",
        "contactphone": "0571-12345678",
        "vendorbanks": [
            {"defaultbank": False, "stopstatus": False, "openaccountbank_name": "非默认银行", "account": "111"},
            {"defaultbank": True, "stopstatus": False, "openaccountbank_name": "默认银行", "account": "222"},
        ],
        "vendorcontactss": [
            {"defaultcontact": True, "contactname": "李四", "contactmobile": "13800000000", "contactemail": "li@example.com"},
        ],
    }]

    with patch("agent.yonyou_vendor.query_supplier_by_name", return_value={"recordCount": 1, "records": records}):
        patch_payload = supplier_patch_from_yonbip(extracted)
    changed = apply_yonbip_supplier_patch(extracted, patch_payload)

    assert patch_payload["matched"] is True
    assert patch_payload["source"] == "yonbip"
    assert patch_payload["missingYonbipFields"] == []
    assert changed == {
        "supplierName",
        "supplierTaxNo",
        "supplierAddress",
        "supplierPhone",
        "supplierBank",
        "supplierAccount",
        "supplierRepresentativeName",
        "supplierRepresentativePhone",
        "supplierRepresentativeEmail",
    }
    assert extracted["supplierName"] == "用友供应商A"
    assert extracted["supplierAddress"] == "用友地址"
    assert extracted["supplierBank"] == "默认银行"
    assert extracted["supplierAccount"] == "222"
    assert extracted["supplierRepresentativeName"] == "李四"
    assert extracted["supplierRepresentativePhone"] == "13800000000"


def test_supplier_patch_from_yonbip_uses_detail_contacts_when_missing_from_page() -> None:
    extracted = {"supplierName": "供应商A"}
    records = [{"id": "vendor-1", "name": "用友供应商A"}]
    detail = {
        "vendorcontactss": [
            {"contactname": "王五", "contactmobile": "13900000000", "contactemail": "wang@example.com"},
        ],
    }

    with patch("agent.yonyou_vendor.query_supplier_by_name", return_value={"recordCount": 1, "records": records}), patch(
        "agent.yonyou_vendor.query_supplier_detail",
        return_value=detail,
    ):
        patch_payload = supplier_patch_from_yonbip(extracted)

    assert patch_payload["patch"]["supplierRepresentativeName"] == "王五"
    assert patch_payload["patch"]["supplierRepresentativePhone"] == "13900000000"
    assert patch_payload["patch"]["supplierRepresentativeEmail"] == "wang@example.com"


def test_supplier_patch_from_yonbip_not_found_and_ambiguous() -> None:
    extracted = {"supplierName": "供应商A"}

    with patch("agent.yonyou_vendor.query_supplier_by_name", return_value={"recordCount": 0, "records": []}):
        not_found = supplier_patch_from_yonbip(extracted)
    with patch(
        "agent.yonyou_vendor.query_supplier_by_name",
        return_value={"recordCount": 2, "records": [{"name": "供应商A"}, {"name": "供应商A"}]},
    ):
        ambiguous = supplier_patch_from_yonbip(extracted)

    assert not_found["matched"] is False
    assert not_found["reason"] == "not_found"
    assert ambiguous["matched"] is False
    assert ambiguous["reason"] == "ambiguous"


def test_parse_dashscope_model_chain_deduplicates() -> None:
    assert contract_llm.parse_model_chain(" qwen3.6-plus ", "qwen-plus, qwen-turbo, qwen-plus, ") == [
        "qwen3.6-plus",
        "qwen-plus",
        "qwen-turbo",
    ]


def test_dashscope_fallback_model_succeeds_after_timeout(monkeypatch) -> None:
    class FakeMessage:
        def __init__(self, content: str) -> None:
            self.content = content

    class FakeChoice:
        def __init__(self, content: str) -> None:
            self.message = FakeMessage(content)

    class FakeCompletion:
        def __init__(self, content: str) -> None:
            self.choices = [FakeChoice(content)]

    class FakeCompletions:
        def __init__(self, responses: dict[str, object], calls: list[str]) -> None:
            self.responses = responses
            self.calls = calls

        def create(self, **kwargs: object) -> FakeCompletion:
            model = str(kwargs["model"])
            self.calls.append(model)
            response = self.responses[model]
            if isinstance(response, Exception):
                raise response
            return FakeCompletion(str(response))

    class FakeChat:
        def __init__(self, responses: dict[str, object], calls: list[str]) -> None:
            self.completions = FakeCompletions(responses, calls)

    class FakeOpenAI:
        calls: list[str] = []
        max_retries: object = None
        responses: dict[str, object] = {
            "primary-model": TimeoutError("Request timed out."),
            "fallback-model": json.dumps({"supplierName": "供应商B", "items": []}, ensure_ascii=False),
        }

        def __init__(self, **kwargs: object) -> None:
            FakeOpenAI.max_retries = kwargs.get("max_retries")
            self.chat = FakeChat(FakeOpenAI.responses, FakeOpenAI.calls)

    monkeypatch.setenv("DASHSCOPE_API_KEY", "key")
    monkeypatch.setenv("DASHSCOPE_MODEL", "primary-model")
    monkeypatch.setenv("DASHSCOPE_FALLBACK_MODELS", "fallback-model,primary-model")
    monkeypatch.delenv("DASHSCOPE_MAX_RETRIES", raising=False)

    with patch("agent.contract.llm.OpenAI", FakeOpenAI):
        result = contract_llm.extract_template_render_data("报价文本", get_template_config("caigouhetong"))

    assert result["supplierName"] == "供应商B"
    assert FakeOpenAI.calls == ["primary-model", "fallback-model"]
    assert FakeOpenAI.max_retries == 0


def test_timeout_error_chain_summary_detects_httpx_layer() -> None:
    FakeReadTimeout = type("ReadTimeout", (Exception,), {"__module__": "httpx"})
    exc = RuntimeError("outer")
    exc.__cause__ = FakeReadTimeout("The read operation timed out")

    assert contract_llm.is_timeout_error(exc)
    assert contract_llm.timeout_layer(exc) == "httpx"
    chain = contract_llm.error_chain_summary(exc)
    assert [item["type"] for item in chain] == ["RuntimeError", "ReadTimeout"]


def test_dashscope_all_models_failed_raises_last_error(monkeypatch) -> None:
    class FakeCompletions:
        calls: list[str] = []

        def create(self, **kwargs: object) -> object:
            model = str(kwargs["model"])
            self.calls.append(model)
            raise TimeoutError(f"{model} timed out")

    class FakeChat:
        def __init__(self) -> None:
            self.completions = FakeCompletions()

    class FakeOpenAI:
        def __init__(self, **_kwargs: object) -> None:
            self.chat = FakeChat()

    monkeypatch.setenv("DASHSCOPE_API_KEY", "key")
    monkeypatch.setenv("DASHSCOPE_MODEL", "primary-model")
    monkeypatch.setenv("DASHSCOPE_FALLBACK_MODELS", "fallback-model")

    with patch("agent.contract.llm.OpenAI", FakeOpenAI), pytest.raises(TimeoutError, match="fallback-model timed out"):
        contract_llm.extract_template_render_data("报价文本", get_template_config("caigouhetong"))

    assert FakeCompletions.calls == ["primary-model", "fallback-model"]


def test_upload_owner_cannot_access_other_user_upload() -> None:
    upload_body = upload_quote()
    response = client.post(
        f"/api/uploads/{upload_body['id']}/quote-text",
        headers=agent_auth_header(userid="uid2", unionid="union-y"),
        json={"templateType": "caigouhetong"},
    )
    assert response.status_code == 403
    assert response.json()["code"] == "FORBIDDEN"


def test_parse_uploaded_quote_image_text() -> None:
    upload = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.png",
            "mimeType": "image/png",
            "size": len(PNG_BYTES),
            "data": base64.b64encode(PNG_BYTES).decode("ascii"),
        },
    )
    assert upload.status_code == 200
    upload_body = upload.json()

    with patch("agent.contract.extract.extract_image_text", return_value="报价单文本"):
        response = client.post(
            f"/api/uploads/{upload_body['id']}/quote-text",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["uploadId"] == upload_body["id"]
    assert body["originalName"] == "quote.png"
    assert body["quoteText"] == "报价单文本"
    assert body["textLength"] == len("报价单文本")
    assert body["parser"] == {"type": "image", "ocrUsed": True}


def test_parse_uploaded_image_uses_ocr() -> None:
    upload = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.png",
            "mimeType": "image/png",
            "size": len(PNG_BYTES),
            "data": base64.b64encode(PNG_BYTES).decode("ascii"),
        },
    )
    assert upload.status_code == 200
    upload_body = upload.json()

    with patch("agent.contract.extract.extract_image_text", return_value="OCR 报价单文本"):
        response = client.post(
            f"/api/uploads/{upload_body['id']}/quote-text",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["quoteText"] == "OCR 报价单文本"
    assert body["parser"] == {"type": "image", "ocrUsed": True}


def test_parse_uploaded_image_ocr_error_returns_ocr_failed() -> None:
    upload_body = upload_quote(original_name="quote.png", mime_type="image/png", content=PNG_BYTES)

    with patch("agent.contract.extract.extract_image_text", side_effect=ValueError("图片 OCR 识别失败：ocrServiceNotOpen")):
        response = client.post(
            f"/api/uploads/{upload_body['id']}/quote-text",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong"},
        )

    assert response.status_code == 502
    body = response.json()
    assert body["code"] == "OCR_FAILED"


def test_extract_image_text_uses_recognize_all_text(tmp_path: Path) -> None:
    from agent.contract.extract import extract_image_text

    image_path = tmp_path / "quote.png"
    image_path.write_bytes(PNG_BYTES)
    calls: dict[str, bool] = {"all_text": False, "general": False}

    class FakeOcrClient:
        def __init__(self, config: object) -> None:
            self.config = config

        def recognize_all_text_with_options(self, request: object, runtime: object) -> object:
            calls["all_text"] = True
            assert getattr(request, "type") == "General"
            assert getattr(request, "body").read().startswith(b"\x89PNG")

            class Response:
                def to_map(self) -> dict:
                    return {"body": {"Data": {"Content": "报价单文本"}}}

            return Response()

        def recognize_general_with_options(self, request: object, runtime: object) -> object:
            calls["general"] = True
            raise AssertionError("should not call RecognizeGeneral")

    class FakeRecognizeAllTextRequest:
        def __init__(self, **kwargs: object) -> None:
            self.__dict__.update(kwargs)

    class FakeConfig:
        def __init__(self, **kwargs: object) -> None:
            self.__dict__.update(kwargs)

    class FakeRuntimeOptions:
        pass

    fake_ocr_root = types.ModuleType("alibabacloud_ocr_api20210707")
    fake_ocr_client = types.ModuleType("alibabacloud_ocr_api20210707.client")
    fake_ocr_models = types.ModuleType("alibabacloud_ocr_api20210707.models")
    fake_openapi = types.ModuleType("alibabacloud_tea_openapi")
    fake_openapi_models = types.ModuleType("alibabacloud_tea_openapi.models")
    fake_util = types.ModuleType("alibabacloud_tea_util")
    fake_util_models = types.ModuleType("alibabacloud_tea_util.models")

    fake_ocr_client.Client = FakeOcrClient
    fake_ocr_models.RecognizeAllTextRequest = FakeRecognizeAllTextRequest
    fake_ocr_root.models = fake_ocr_models
    fake_openapi.models = fake_openapi_models
    fake_openapi_models.Config = FakeConfig
    fake_util.models = fake_util_models
    fake_util_models.RuntimeOptions = FakeRuntimeOptions

    with patch.dict(
        sys.modules,
        {
            "alibabacloud_ocr_api20210707": fake_ocr_root,
            "alibabacloud_ocr_api20210707.client": fake_ocr_client,
            "alibabacloud_ocr_api20210707.models": fake_ocr_models,
            "alibabacloud_tea_openapi": fake_openapi,
            "alibabacloud_tea_openapi.models": fake_openapi_models,
            "alibabacloud_tea_util": fake_util,
            "alibabacloud_tea_util.models": fake_util_models,
        },
    ), patch.dict(os.environ, {"ALIYUN_ACCESS_KEY_ID": "ak", "ALIYUN_ACCESS_KEY_SECRET": "sk"}):
        assert extract_image_text(image_path) == "报价单文本"

    assert calls == {"all_text": True, "general": False}


def test_upload_rejects_empty_file() -> None:
    response = client.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "empty.pdf",
            "mimeType": "application/pdf",
            "size": 0,
            "data": "",
        },
    )
    assert response.status_code == 400
    assert "文件内容" in response.json()["message"]


def test_contract_generate_uses_confirmed_quote_text() -> None:
    upload_id = upload_quote(content=b"%PDF-1.4 raw quote")["id"]

    with patch("agent.main.generate_contract") as generate_contract:
        generate_contract.return_value = {
            "contractId": "contract_test",
            "templateType": "caigouhetong",
            "quoteTextLength": 4,
            "dingDrive": {"spaceId": "space1", "fileId": "file1", "fileName": "confirmed.docx"},
        }
        response = client.post(
            "/api/contracts/generate",
            headers=agent_auth_header(),
            json={
                "uploadId": upload_id,
                "templateType": "caigouhetong",
                "quoteText": " 用户确认文本 ",
            },
        )

    assert response.status_code == 200
    generate_contract.assert_called_once_with(upload_id, "caigouhetong", "用户确认文本", None, None, ANY, "auto")
    assert response.json()["dingDrive"]["fileId"] == "file1"


def test_field_preview_uses_extra_info_and_classifies_fields() -> None:
    upload_id = upload_quote()["id"]

    extracted = {
        "supplierName": "供应商A",
        "buyerPhone": None,
        "items": [{"index": "1", "name": "水泵", "quantity": "2", "unitPrice": "100"}],
    }
    with patch("agent.main.extract_template_render_data", return_value=extracted) as llm:
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={
                "templateType": "caigouhetong",
                "quoteText": " 用户确认报价单文本 ",
                "extraInfo": " 付款方式：验收后支付 ",
            },
        )

    assert response.status_code == 200
    llm.assert_called_once_with("用户确认报价单文本", ANY, "付款方式：验收后支付")
    body = response.json()
    assert body["extractedData"] == extracted
    assert any(field["label"] == "乙方名称" and field["value"] == "供应商A" for field in body["recognizedFields"])
    assert any(field["label"] == "甲方电话" for field in body["missingFields"])
    assert body["tableRowCounts"] == {"items": 1}


def test_field_preview_applies_yonbip_supplier_patch() -> None:
    upload_id = upload_quote()["id"]
    extracted = {"supplierName": "供应商A", "supplierAddress": "报价单地址", "supplierBank": "报价单开户行", "items": []}

    def fake_supplier_patch(data: dict) -> dict:
        return {
            "source": "yonbip",
            "matched": True,
            "patch": {"supplierAddress": "用友地址", "supplierBank": "用友开户行"},
            "missingYonbipFields": [],
        }

    with patch("agent.main.extract_template_render_data", return_value=extracted), patch(
        "agent.main.supplier_patch_from_yonbip",
        side_effect=fake_supplier_patch,
    ):
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong", "quoteText": "用户确认报价单文本"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["extractedData"]["supplierAddress"] == "用友地址"
    assert body["extractedData"]["supplierBank"] == "用友开户行"
    assert body["supplierPatch"]["source"] == "yonbip"
    assert body["supplierPatch"]["overwrittenFields"] == ["supplierAddress", "supplierBank"]


def test_field_preview_yonbip_failure_does_not_block() -> None:
    upload_id = upload_quote()["id"]

    with patch("agent.main.extract_template_render_data", return_value={"supplierName": "供应商A", "items": []}), patch(
        "agent.main.supplier_patch_from_yonbip",
        side_effect=RuntimeError("YonBIP timeout"),
    ):
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong", "quoteText": "用户确认报价单文本"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["supplierPatch"]["source"] == "yonbip"
    assert body["supplierPatch"]["matched"] is False
    assert body["supplierPatch"]["reason"] == "lookup_error"


def test_supplier_lookup_api_success() -> None:
    fake_patch = {
        "source": "yonbip",
        "matched": True,
        "patch": {"supplierTaxNo": "9133", "supplierAccount": "6222"},
        "missingYonbipFields": ["supplierBank"],
    }
    with patch("agent.main.supplier_patch_from_yonbip", return_value=fake_patch):
        response = client.post(
            "/api/suppliers/lookup",
            headers=agent_auth_header(),
            json={"supplierName": "供应商A"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["ok"] is True
    assert body["supplierPatch"]["matched"] is True
    assert body["supplierPatch"]["patch"]["supplierTaxNo"] == "9133"
    assert body["supplierPatch"]["missingYonbipFields"] == ["supplierBank"]


def test_supplier_lookup_api_missing_name() -> None:
    response = client.post(
        "/api/suppliers/lookup",
        headers=agent_auth_header(),
        json={"supplierName": "  "},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["supplierPatch"]["matched"] is False
    assert body["supplierPatch"]["reason"] == "missing_supplier_name"


def test_supplier_lookup_api_yonbip_error() -> None:
    with patch("agent.main.supplier_patch_from_yonbip", side_effect=RuntimeError("YonBIP timeout")):
        response = client.post(
            "/api/suppliers/lookup",
            headers=agent_auth_header(),
            json={"supplierName": "供应商A"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["supplierPatch"]["matched"] is False
    assert body["supplierPatch"]["reason"] == "lookup_error"


def test_field_preview_complex_excel_uses_scalar_only_contract() -> None:
    content = xlsx_bytes([["品名", "数量"], *[[f"设备{i}", i] for i in range(1, 6)]])
    upload_id = upload_quote(
        original_name="quote.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=content,
    )["id"]

    def fake_llm(_quote_text: str, config: object, _extra_info: str | None = None) -> dict:
        assert getattr(config, "table_bindings") == {}
        return {"supplierName": "供应商A"}

    with patch("agent.main.extract_template_render_data", side_effect=fake_llm), patch(
        "agent.main.supplier_patch_from_yonbip",
        return_value={"source": "yonbip", "matched": False, "patch": {}, "reason": "not_found"},
    ):
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong", "quoteText": "用户确认报价单文本"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["attachmentMode"]["enabled"] is True
    assert body["tableRowCounts"] == {}
    assert not any(field["type"] == "table" for field in body["recognizedFields"])
    assert not any(field["type"] in {"table", "tableCell"} for field in body["missingFields"])


def test_field_preview_template_table_mode_keeps_table_contract() -> None:
    content = xlsx_bytes([["品名", "数量"], *[[f"设备{i}", i] for i in range(1, 6)]])
    upload_id = upload_quote(
        original_name="quote.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=content,
    )["id"]

    def fake_llm(_quote_text: str, config: object, _extra_info: str | None = None) -> dict:
        assert getattr(config, "table_bindings") != {}
        return {"supplierName": "供应商A", "items": [{"name": "设备1", "quantity": "1"}]}

    with patch("agent.main.extract_template_render_data", side_effect=fake_llm), patch(
        "agent.main.supplier_patch_from_yonbip",
        return_value={"source": "yonbip", "matched": False, "patch": {}, "reason": "not_found"},
    ):
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={"templateType": "caigouhetong", "quoteText": "用户确认报价单文本", "tableMode": "template"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["attachmentMode"]["enabled"] is False
    assert body["tableMode"] == "template"
    assert body["tableRowCounts"]["items"] == 1


def test_field_preview_timeout_returns_retryable_message() -> None:
    upload_id = upload_quote()["id"]

    with patch("agent.main.extract_template_render_data", side_effect=TimeoutError("Request timed out.")):
        response = client.post(
            f"/api/uploads/{upload_id}/field-preview",
            headers=agent_auth_header(),
            json={
                "templateType": "caigouhetong",
                "quoteText": "用户确认报价单文本",
            },
        )

    assert response.status_code == 502
    body = response.json()
    assert body["code"] == "LLM_FAILED"
    assert body["message"] == "字段识别超时，请稍后重试；如报价单内容较长，可先删减无关文本后再识别"
    assert "Traceback" not in json.dumps(body, ensure_ascii=False)


def test_generate_contract_reuses_confirmed_extracted_data() -> None:
    upload_id = upload_quote()["id"]
    extracted = {"supplierName": "供应商A", "items": []}

    with patch("agent.main.extract_template_render_data") as llm, patch(
        "agent.main.render_contract",
        return_value=Path("agent/storage/contracts/confirmed.docx"),
    ) as render_contract_mock, patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "confirmed.docx", "filePath": "/confirmed.docx"},
    ):
        draft = generate_contract(upload_id, "caigouhetong", "确认文本", "补充信息", extracted, {"userid": "uid1", "unionid": "union-x"})

    llm.assert_not_called()
    render_contract_mock.assert_called_once_with(
        ANY,
        ANY,
        ANY,
        blank_missing=True,
        table_mode="auto",
        quote_attachment=None,
        logger=ANY,
    )
    assert draft["extractedData"] == extracted
    assert draft["extraInfoLength"] == len("补充信息")
    assert "supplierCacheWriteback" not in draft
    assert draft["dingDrive"]["fileId"] == "file1"


def test_generate_contract_attachment_table_mode_passes_quote_attachment() -> None:
    content = xlsx_bytes([["品名", "数量"], ["设备1", 1]])
    upload_id = upload_quote(
        original_name="quote.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=content,
    )["id"]
    extracted = {"supplierName": "供应商A", "items": []}

    with patch("agent.main.extract_template_render_data") as llm, patch(
        "agent.main.render_contract",
        return_value=Path("agent/storage/contracts/attachment.docx"),
    ) as render_contract_mock, patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "attachment.docx", "filePath": "/attachment.docx"},
    ):
        draft = generate_contract(
            upload_id,
            "caigouhetong",
            "确认文本",
            "补充信息",
            extracted,
            {"userid": "uid1", "unionid": "union-x"},
            table_mode="attachment",
        )

    llm.assert_not_called()
    quote_attachment = render_contract_mock.call_args.kwargs["quote_attachment"]
    assert quote_attachment["attachmentMode"]["enabled"] is True
    assert quote_attachment["attachmentMode"]["tableMode"] == "attachment"
    assert draft["attachmentMode"]["enabled"] is True
    assert draft["tableMode"] == "attachment"


def test_contract_file_name_uses_contract_supplier_and_project() -> None:
    assert contract_file_name({
        "contractNo": "HT/001",
        "supplierName": "供应商A",
        "projectName": "项目:一",
    }) == "HT_001_供应商A_项目_一.docx"
    assert contract_file_name({}, generated_at=date(2026, 6, 1)) == "20260601_000000_未知乙方_未知项目.docx"


def test_tax_fields_are_calculated_from_total_amount() -> None:
    config = get_template_config("caigouhetong")
    extracted = {"totalAmount": "113", "taxRate": "13"}

    changed = apply_tax_calculations(extracted, config)

    assert changed == {"amountWithoutTax", "taxAmount"}
    assert extracted["amountWithoutTax"] == "100"
    assert extracted["taxAmount"] == "13"


def test_line_item_total_is_calculated_from_quantity_and_unit_price() -> None:
    config = get_template_config("simpleContract")
    extracted = {
        "items": [
            {"index": "1", "name": "产品A", "quantity": "2", "unitPrice": "10", "totalPrice": ""},
        ],
    }

    changed = apply_line_item_calculations(extracted, config)

    assert changed == {"items.totalPrice"}
    assert extracted["items"][0]["totalPrice"] == "20"


def test_line_item_total_skips_when_quantity_or_unit_price_missing() -> None:
    config = get_template_config("simpleContract")
    extracted = {
        "items": [
            {"index": "1", "name": "产品A", "quantity": "2", "unitPrice": "", "totalPrice": ""},
        ],
    }

    changed = apply_line_item_calculations(extracted, config)

    assert changed == set()
    assert extracted["items"][0]["totalPrice"] == ""


def test_generate_contract_without_upload_id_uses_confirmed_data() -> None:
    extracted = {
        "supplierName": "供应商A",
        "items": [{"index": "1", "name": "阀门", "quantity": "2", "unitPrice": "10", "totalPrice": "20"}],
    }

    with patch("agent.main.extract_template_render_data") as llm, patch(
        "agent.main.render_contract",
        return_value=Path("agent/storage/contracts/manual.docx"),
    ) as render_contract_mock, patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "manual.docx", "filePath": "/manual.docx"},
    ), patch(
        "agent.main.remove_upload",
    ) as remove_upload_mock:
        draft = generate_contract(None, "caigouhetong", None, None, extracted, {"userid": "uid1", "unionid": "union-x"})

    llm.assert_not_called()
    remove_upload_mock.assert_not_called()
    assert render_contract_mock.call_args.kwargs["table_mode"] == "template"
    assert render_contract_mock.call_args.kwargs["quote_attachment"] is None
    assert draft["extractedData"]["items"][0]["totalPrice"] == "20"
    assert draft["upload"] is None


def test_delivery_date_is_calculated_from_delivery_days() -> None:
    config = get_template_config("caigouhetong")
    extracted = {"deliveryDays": "7"}

    changed = apply_delivery_date_calculation(extracted, config, today=date(2026, 5, 28))

    assert changed == {"deliveryYear", "deliveryMonth", "deliveryDay"}
    assert extracted["deliveryYear"] == "2026"
    assert extracted["deliveryMonth"] == "06"
    assert extracted["deliveryDay"] == "04"


def test_delivery_date_does_not_override_confirmed_date() -> None:
    config = get_template_config("caigouhetong")
    extracted = {"deliveryDays": "7", "deliveryYear": "2026", "deliveryMonth": "12", "deliveryDay": "31"}

    changed = apply_delivery_date_calculation(extracted, config, today=date(2026, 5, 28))

    assert changed == set()
    assert extracted["deliveryYear"] == "2026"
    assert extracted["deliveryMonth"] == "12"
    assert extracted["deliveryDay"] == "31"


def test_generate_contract_backfills_confirmed_tax_fields() -> None:
    upload_id = upload_quote()["id"]
    extracted = {"supplierName": "供应商A", "totalAmount": "113", "taxRate": "13", "items": []}

    with patch("agent.main.extract_template_render_data") as llm, patch(
        "agent.main.render_contract",
        return_value=Path("agent/storage/contracts/tax.docx"),
    ) as render_contract_mock, patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "tax.docx", "filePath": "/tax.docx"},
    ):
        draft = generate_contract(upload_id, "caigouhetong", "确认文本", "补充信息", extracted, {"userid": "uid1", "unionid": "union-x"})

    llm.assert_not_called()
    render_data = render_contract_mock.call_args.args[0]
    assert render_data["amountWithoutTax"] == "100"
    assert render_data["taxAmount"] == "13"
    assert draft["extractedData"]["amountWithoutTax"] == "100"
    assert draft["extractedData"]["taxAmount"] == "13"


def test_generate_contract_backfills_confirmed_delivery_date() -> None:
    upload_id = upload_quote()["id"]
    extracted = {"supplierName": "供应商A", "deliveryDays": "7", "items": []}

    with patch("agent.main.extract_template_render_data") as llm, patch(
        "agent.main._today_shanghai",
        return_value=date(2026, 5, 28),
    ), patch(
        "agent.main.render_contract",
        return_value=Path("agent/storage/contracts/delivery.docx"),
    ) as render_contract_mock, patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "delivery.docx", "filePath": "/delivery.docx"},
    ):
        draft = generate_contract(upload_id, "caigouhetong", "确认文本", "补充信息", extracted, {"userid": "uid1", "unionid": "union-x"})

    llm.assert_not_called()
    render_data = render_contract_mock.call_args.args[0]
    assert render_data["deliveryYear"] == "2026"
    assert render_data["deliveryMonth"] == "06"
    assert render_data["deliveryDay"] == "04"
    assert draft["extractedData"]["deliveryYear"] == "2026"


def test_generate_contract_uploads_dingdrive_and_removes_process_files() -> None:
    upload_body = upload_quote()
    upload_id = upload_body["id"]
    record = upload_record(upload_id)
    extracted = {"supplierName": "供应商A", "items": []}
    rendered_path = Path("agent/storage/contracts/20260523_供应商A.docx")

    def fake_render(*args, **kwargs) -> Path:
        rendered_path.parent.mkdir(parents=True, exist_ok=True)
        rendered_path.write_bytes(b"docx")
        return rendered_path

    with patch("agent.main.render_contract", side_effect=fake_render), patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "20260523_供应商A.docx", "filePath": "/采购合同测试/20260523_供应商A.docx"},
    ) as upload_dingdrive:
        draft = generate_contract(upload_id, "caigouhetong", "确认文本", "补充信息", extracted, {"userid": "uid1", "unionid": "union-x"})

    upload_dingdrive.assert_called_once()
    assert draft["dingDrive"]["fileId"] == "file1"
    assert "供应商A" in draft["fileName"]
    assert not Path(record["path"]).exists()
    assert not (UPLOADS_DIR / f"{upload_id}.json").exists()
    assert not rendered_path.exists()
    assert not list(UPLOADS_DIR.parent.glob(f"**/{draft['contractId']}.json"))


def test_contract_download_payload_returns_dingdrive_download_info() -> None:
    draft = {
        "contractId": "contract_test",
        "fileName": "20260523_供应商A.docx",
        "dingDrive": {
            "spaceId": "space1",
            "fileId": "file1",
            "fileName": "20260523_供应商A.docx",
            "filePath": "/采购合同测试/20260523_供应商A.docx",
            "fileSize": 1234,
            "fileType": "docx",
        },
    }

    payload = contract_download_payload(draft)

    assert payload["download"]["type"] == "agent_proxy"
    assert payload["download"]["fileName"] == "20260523_供应商A.docx"
    assert payload["fileSize"] == 1234
    assert payload["fileType"] == "docx"


def test_dingdrive_download_proxy_streams_file() -> None:
    upstream = Mock()
    upstream.iter_content.return_value = [b"docx-content"]
    upstream.raise_for_status.return_value = None

    with patch(
        "agent.main.get_contract_download_info",
        return_value={"resourceUrls": ["https://download.example/file"], "headers": {"x-test": "1"}},
    ) as download_info, patch("agent.main.requests.get", return_value=upstream) as get:
        response = client.post(
            "/api/dingdrive/download",
            headers=agent_auth_header(),
            json={"spaceId": "space1", "fileId": "file1", "fileName": "20260523_供应商A.docx"},
        )

    assert response.status_code == 200
    assert response.content == b"docx-content"
    assert "filename*=UTF-8''20260523_" in response.headers["content-disposition"]
    download_info.assert_called_once()
    get.assert_called_once_with("https://download.example/file", headers={"x-test": "1"}, stream=True, timeout=120)
    upstream.close.assert_called_once()


def test_generate_contract_keeps_process_files_when_dingdrive_fails() -> None:
    upload_body = upload_quote()
    upload_id = upload_body["id"]
    record = upload_record(upload_id)
    rendered_path = Path("agent/storage/contracts/failed-upload.docx")

    def fake_render(*args, **kwargs) -> Path:
        rendered_path.parent.mkdir(parents=True, exist_ok=True)
        rendered_path.write_bytes(b"docx")
        return rendered_path

    with patch("agent.main.render_contract", side_effect=fake_render), patch(
        "agent.main.upload_contract_to_dingdrive",
        side_effect=RuntimeError("钉盘上传失败"),
    ), pytest.raises(RuntimeError):
        generate_contract(upload_id, "caigouhetong", "确认文本", "补充信息", {"supplierName": "供应商A", "items": []}, {"userid": "uid1", "unionid": "union-x"})

    assert Path(record["path"]).exists()
    assert (UPLOADS_DIR / f"{upload_id}.json").exists()
    assert rendered_path.exists()
    rendered_path.unlink(missing_ok=True)


def test_confirmed_blank_fields_render_empty() -> None:
    config = get_template_config("caigouhetong")
    render_data = {"projectName": "", "supplierName": "", "items": [{"index": "1", "name": ""}]}

    pending_context = build_docxtpl_context(render_data, config)
    confirmed_context = build_docxtpl_context(render_data, config, blank_missing=True)

    assert "【待填写：项目名称】" in str(pending_context["projectName"])
    assert "【待填写：乙方名称】" in str(pending_context["supplierName"])
    project_xml = str(confirmed_context["projectName"])
    supplier_xml = str(confirmed_context["supplierName"])
    item_xml = str(confirmed_context["items"][0]["name"])
    assert 'w:eastAsia="仿宋"' in project_xml
    assert 'w:sz w:val="21"' in project_xml
    assert '<w:u w:val="single"/>' in project_xml
    assert "        " in project_xml
    assert '<w:u w:val="single"/>' not in supplier_xml
    assert '<w:u w:val="single"/>' not in item_xml


def test_labor_subcontract_context_uses_fangsong_xiaosi() -> None:
    config = get_template_config("laborSubcontract")
    context = build_docxtpl_context({"projectName": "测试项目", "items": []}, config, blank_missing=True)
    project_xml = str(context["projectName"])
    assert 'w:eastAsia="仿宋"' in project_xml
    assert 'w:sz w:val="24"' in project_xml


def test_simple_contract_context_uses_songti() -> None:
    config = get_template_config("simpleContract")
    context = build_docxtpl_context({"contractNo": "SC-001", "items": []}, config, blank_missing=True)
    contract_xml = str(context["contractNo"])
    assert 'w:eastAsia="宋体"' in contract_xml
    assert 'w:sz w:val="21"' in contract_xml


def test_append_quote_attachment_adds_excel_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract.docx"
    Document().save(docx_path)
    typography = get_template_typography("caigouhetong")

    append_quote_attachment(docx_path, {
        "sheets": [{
            "name": "报价",
            "rows": [["品名", "数量"], ["阀门", "2"]],
        }],
    }, typography)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "附件：报价单明细" in paragraph_text
    assert "报价" in paragraph_text
    assert "品名" in table_text
    assert "阀门" in table_text


def test_append_quote_attachment_skips_empty_rows_without_adding_page(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract.docx"
    Document().save(docx_path)
    original = Document(docx_path)
    original_paragraph_count = len(original.paragraphs)

    append_quote_attachment(docx_path, {
        "sheets": [{
            "name": "报价",
            "rows": [["", ""], ["", ""]],
        }],
    }, get_template_typography("caigouhetong"))

    document = Document(docx_path)
    assert len(document.paragraphs) == original_paragraph_count
    assert len(document.tables) == 0
    assert all(paragraph.text != "附件：报价单明细" for paragraph in document.paragraphs)


def test_append_quote_attachment_does_not_require_heading_styles(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract.docx"
    Document().save(docx_path)

    append_quote_attachment(docx_path, {
        "sheets": [{
            "name": "报价",
            "rows": [["品名", "数量"], ["阀门", "2"]],
        }],
    }, get_template_typography("caigouhetong"))

    document = Document(docx_path)
    assert any(paragraph.text == "附件：报价单明细" for paragraph in document.paragraphs)
    assert any(paragraph.text == "报价" for paragraph in document.paragraphs)


def test_append_quote_attachment_writes_table_borders_without_style(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract.docx"
    Document().save(docx_path)

    append_quote_attachment(docx_path, {
        "sheets": [{
            "name": "报价",
            "rows": [["品名", "数量"], ["阀门", "2"]],
        }],
    }, get_template_typography("caigouhetong"))

    with ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    assert "<w:tblBorders>" in xml
    assert '<w:tblLayout w:type="fixed"' in xml
    assert '<w:tblW w:w="5000" w:type="pct"' in xml
    assert '<w:insideH w:val="single"' in xml
    assert '<w:insideV w:val="single"' in xml


def _grid_col_widths_from_xml(xml: str, expected_cols: int | None = None) -> list[int]:
    widths = [int(width) for width in re.findall(r'<w:gridCol[^>]*w:w="(\d+)"', xml)]
    if expected_cols is not None:
        assert len(widths) >= expected_cols, f"document.xml 中 gridCol 数量不足：{len(widths)} < {expected_cols}"
        return widths[-expected_cols:]
    tbl_blocks = re.findall(r"<w:tbl>.*?</w:tbl>", xml, re.DOTALL)
    assert tbl_blocks, "document.xml 中未找到表格"
    tbl_xml = tbl_blocks[-1]
    return [int(width) for width in re.findall(r'<w:gridCol w:w="(\d+)"', tbl_xml)]


def test_choose_attachment_write_strategy_tiers() -> None:
    assert choose_attachment_write_strategy(2, 2) == "xml_bulk"
    assert choose_attachment_write_strategy(10, 20) == "xml_bulk"
    assert choose_attachment_write_strategy(100, 8) == "xml_bulk"
    assert choose_attachment_write_strategy(451, 10) == "xml_bulk"
    assert choose_attachment_write_strategy(500, 50) == "xml_chunked"


def _attachment_rows(row_count: int, col_count: int) -> list[list[str]]:
    header = [f"列{index}" for index in range(1, col_count + 1)]
    body = [[f"行{row}列{col}" for col in range(1, col_count + 1)] for row in range(1, row_count)]
    return [header, *body]


def test_append_quote_attachment_large_table_uses_xml_bulk(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract-large.docx"
    Document().save(docx_path)
    rows = _attachment_rows(451, 10)

    append_quote_attachment(docx_path, {"sheets": [{"name": "报价", "rows": rows}]}, get_template_typography("caigouhetong"))

    document = Document(docx_path)
    table = document.tables[-1]
    assert len(table.rows) == 451
    assert len(table.rows[0].cells) == 10
    assert table.cell(0, 0).text == "列1"
    assert table.cell(450, 9).text == "行450列10"
    with ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    assert "<w:tblBorders>" in xml
    assert '<w:tblW w:w="5000" w:type="pct"' in xml
    assert '<w:tblLayout w:type="fixed"' in xml
    grid_widths = _grid_col_widths_from_xml(xml, expected_cols=10)
    assert len(grid_widths) == 10
    assert sum(grid_widths) == ATTACHMENT_TABLE_CONTENT_WIDTH_DXA
    assert 'w:orient="landscape"' in xml


def test_column_display_weights_use_max_header_body() -> None:
    assert _text_display_width("单价 (元)") > _text_display_width("1")
    assert _text_display_width("设备名称") > _text_display_width("1")

    rows = [
        ["序号", "设备名称", "单价 (元)"],
        ["1", "控制柜", "1"],
    ]
    widths = compute_attachment_column_widths(rows, 3)
    assert len(widths) == 3
    assert sum(widths) == ATTACHMENT_TABLE_CONTENT_WIDTH_DXA
    assert widths[1] > widths[0]
    assert widths[2] > widths[0]


def test_append_quote_attachment_xml_bulk_column_widths_follow_content(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract-quote-widths.docx"
    Document().save(docx_path)
    header = ["序号", "设备名称", "设备型号", "数量", "单价 (元)", "小计 (元)", "备注"]
    body = [
        ["1", "控制柜主体很长的名称", "X-100", "1", "1", "1", ""],
        ["2", "短名", "Y-200", "2", "2", "2", "详见清单"],
    ]
    body.extend(
        [str(index), "短名", f"Y-{index}", "1", "1", "1", ""]
        for index in range(3, 32)
    )
    rows = [header, *body]

    append_quote_attachment(docx_path, {"sheets": [{"name": "方案报价表", "rows": rows}]}, get_template_typography("caigouhetong"))

    with ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    grid_widths = _grid_col_widths_from_xml(xml, expected_cols=7)
    assert len(grid_widths) == 7
    assert sum(grid_widths) == ATTACHMENT_TABLE_CONTENT_WIDTH_DXA
    assert 'w:orient="landscape"' in xml
    assert grid_widths[1] > grid_widths[0]
    assert grid_widths[2] > grid_widths[0]
    assert grid_widths[4] > grid_widths[3]
    assert grid_widths[5] >= grid_widths[4]
    assert len(set(grid_widths)) > 1


def test_append_quote_attachment_wide_table_spec_column_wider_than_quantity(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract-wide-spec.docx"
    Document().save(docx_path)
    long_spec = (
        "电磁流量计：1、安装位置：室内管道安装；2、技术参数：Q=0-5m3/h，DN25，一体式，法兰连接；"
        "3、材质：哈氏合金电极，PTFE衬里"
    )
    long_param = "【电磁流量计】公称通径：DN25(1\")；过程连接标准：JB/T 81 法兰；公称压力：PN16"
    header = [
        "物料名称",
        "设备位号",
        "设备工艺名称",
        "物料用区",
        "下单规格说明",
        "返资参数",
        "型号",
        "主数量",
        "主单位",
        "含税单价",
        "含税总价",
        "货期",
    ]
    body = [[
        "流量仪表",
        "FIT-105",
        "氨氮废水中转池电磁流量计",
        "废水",
        long_spec,
        long_param,
        "LDG-SUP-A100-25JCMCKAAMGN6WA001-HZD",
        "1",
        "台",
        "1",
        "1",
        "",
    ]]
    body.extend([
        [
            "流量仪表",
            f"FIT-{index}",
            "氨氮废水中转池电磁流量计",
            "废水",
            long_spec,
            long_param,
            "LDG-SUP-A100-25JCMCKAAMGN6WA001-HZD",
            "1",
            "台",
            "1",
            "1",
            "",
        ]
        for index in range(106, 125)
    ])
    rows = [header, *body]

    append_quote_attachment(docx_path, {"sheets": [{"name": "报价", "rows": rows}]}, get_template_typography("caigouhetong"))

    with ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    grid_widths = _grid_col_widths_from_xml(xml, expected_cols=12)
    assert sum(grid_widths) == ATTACHMENT_TABLE_CONTENT_WIDTH_DXA
    assert 'w:orient="landscape"' in xml
    spec_index = header.index("下单规格说明")
    quantity_index = header.index("主数量")
    assert grid_widths[spec_index] > grid_widths[quantity_index] * 5


def test_append_quote_attachment_chunked_table_writes_all_rows(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract-chunked.docx"
    Document().save(docx_path)
    rows = _attachment_rows(500, 50)

    append_quote_attachment(docx_path, {"sheets": [{"name": "超大表", "rows": rows}]}, get_template_typography("caigouhetong"))

    document = Document(docx_path)
    table = document.tables[-1]
    assert len(table.rows) == 500
    assert len(table.rows[0].cells) == 50
    assert table.cell(499, 49).text == "行499列50"
    with ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    assert '<w:tblW w:w="5000" w:type="pct"' in xml
    assert '<w:tblLayout w:type="fixed"' in xml
    grid_widths = _grid_col_widths_from_xml(xml, expected_cols=50)
    assert len(grid_widths) == 50
    assert sum(grid_widths) == ATTACHMENT_TABLE_CONTENT_WIDTH_DXA
    assert 'w:orient="landscape"' in xml


def test_append_quote_attachment_logs_write_strategy(tmp_path: Path) -> None:
    docx_path = tmp_path / "contract-log.docx"
    Document().save(docx_path)
    logs: list[tuple[str, dict[str, Any]]] = []

    def logger(message: str, **meta: Any) -> None:
        logs.append((message, meta))

    append_quote_attachment(
        docx_path,
        {"sheets": [{"name": "报价", "rows": _attachment_rows(100, 8)}]},
        get_template_typography("caigouhetong"),
        logger=logger,
    )

    assert logs
    message, meta = logs[-1]
    assert message == "quote attachment appended"
    assert meta["sheetStats"][0]["writeStrategy"] == "xml_bulk"
    assert meta["sheetStats"][0]["rowCount"] == 100
    assert meta["sheetStats"][0]["colCount"] == 8
    assert meta["elapsedMs"] >= 0


def test_generate_contract_large_attachment_completes_quickly() -> None:
    import time

    rows = _attachment_rows(451, 10)
    content = xlsx_bytes(rows)
    upload_id = upload_quote(
        original_name="large-quote.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=content,
    )["id"]
    extracted = {"supplierName": "浙江沃辉"}
    append_elapsed_holder: dict[str, float] = {}

    def fake_render(_render_data: dict[str, Any], _config: object, contract_id: str, **kwargs: Any) -> Path:
        from agent.contract.config import CONTRACTS_DIR

        path = CONTRACTS_DIR / f"{contract_id}.docx"
        Document().save(path)
        quote_attachment = kwargs.get("quote_attachment")
        assert isinstance(quote_attachment, dict)
        sheet_rows = quote_attachment["sheets"][0]["rows"]
        assert len(sheet_rows) >= 450
        append_start = time.perf_counter()
        append_quote_attachment(path, quote_attachment, get_template_typography("simpleContract"))
        append_elapsed_holder["seconds"] = time.perf_counter() - append_start
        return path

    with patch("agent.main.render_contract", side_effect=fake_render), patch(
        "agent.main.upload_contract_to_dingdrive",
        return_value={"spaceId": "space1", "fileId": "file1", "fileName": "large.docx", "filePath": "/large.docx"},
    ), patch("agent.main.remove_upload", return_value=[]), patch(
        "agent.main.remove_contract_files",
        return_value=[],
    ):
        start = time.perf_counter()
        draft = generate_contract(
            upload_id,
            "simpleContract",
            "用户确认报价单文本",
            None,
            extracted,
            {"userid": "uid1", "unionid": "union-x"},
            table_mode="attachment",
        )
        total_elapsed = time.perf_counter() - start

    assert draft["tableMode"] == "attachment"
    assert draft["attachmentMode"]["enabled"] is True
    assert total_elapsed < 60
    assert append_elapsed_holder["seconds"] < 30


def test_render_contract_centers_template_table_text() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "contractNo": "HT-001",
        "supplierName": "供应商A",
        "projectName": "项目A",
        "items": [{"index": "1", "name": "阀门", "spec": "DN50", "unit": "台", "quantity": "2", "unitPrice": "10", "totalPrice": "20", "tagNo": "T1"}],
    }, config)

    output_path = render_contract(render_data, config, "test_table_format")
    try:
        document = Document(output_path)
        header_tables = document.tables[:2]
        assert len(header_tables) == 2
        for table in header_tables:
            header_paragraphs = [
                paragraph
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            ]
            assert header_paragraphs
            assert all(paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT for paragraph in header_paragraphs)
        paragraphs = [
            paragraph
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            if "阀门" in paragraph.text
        ]
        assert paragraphs
        paragraph = paragraphs[0]
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        formatted_runs = [run for run in paragraph.runs if run.text.strip()]
        assert formatted_runs
        assert formatted_runs[0].font.name == "仿宋"
        assert formatted_runs[0].font.size.pt == 10.5
    finally:
        output_path.unlink(missing_ok=True)


def test_render_contract_body_paragraph_font_after_placeholder() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "contractNo": "HT-001",
        "supplierName": "供应商A",
        "projectName": "项目A",
        "purchaseSubject": "阀门设备",
        "deliveryPlace": "杭州余杭",
        "items": [{"index": "1", "name": "阀门", "spec": "DN50", "unit": "台", "quantity": "2", "unitPrice": "10", "totalPrice": "20", "tagNo": "T1"}],
    }, config)

    output_path = render_contract(render_data, config, "test_body_paragraph_font")
    try:
        document = Document(output_path)
        target_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if "交货地点" in paragraph.text or "就甲方" in paragraph.text
        ]
        assert target_paragraphs
        for paragraph in target_paragraphs:
            formatted_runs = [run for run in paragraph.runs if run.text.strip()]
            assert formatted_runs
            for run in formatted_runs:
                assert run.font.name == "仿宋"
                assert run.font.size.pt == 10.5
    finally:
        output_path.unlink(missing_ok=True)


def test_contract_render_imports_with_fc_layout() -> None:
    import subprocess

    agent_dir = Path(__file__).resolve().parents[1]
    result = subprocess.run(
        [
            sys.executable,
            "-c",
            "import importlib; importlib.import_module('contract.render')",
        ],
        cwd=str(agent_dir),
        env={**os.environ, "PYTHONPATH": str(agent_dir)},
        capture_output=True,
        text=True,
        check=False,
    )
    assert result.returncode == 0, result.stderr or result.stdout


def test_render_annual_framework_extended_scalars() -> None:
    config = get_template_config("annualFramework")
    render_data = merge_render_data({
        "contractNo": "AF-001",
        "purchaseSubject": "阀门及配件",
        "contractTermYears": "1",
        "contractStartDate": "2026年01月01日",
        "contractEndDate": "2026年12月31日",
        "deliveryDays": "3",
        "transportMethod": "快递",
        "shipmentNoticeHours": "8",
        "settlementCycle": "月结30天",
        "supplierName": "年度供应商",
        "supplierAddress": "杭州余杭",
        "supplierAccount": "622200000000",
        "supplierTaxNo": "91330000TEST",
        "supplierPhone": "0571-12345678",
        "supplierFax": "0571-87654321",
        "supplierRepresentativeName": "李四",
        "supplierRepresentativePhone": "13800001111",
        "supplierRepresentativeAddress": "杭州市余杭区仓前街道",
        "supplierRepresentativeEmail": "lisi@example.com",
        "supplierAuthorizedRepresentative": "王五",
        "signatureYear": "2026",
        "signatureMonth": "06",
        "signatureDay": "05",
    }, config)

    output_path = render_contract(render_data, config, "test_annual_framework_extended_scalars", blank_missing=True)
    try:
        with ZipFile(output_path) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
    finally:
        output_path.unlink(missing_ok=True)

    assert "0571-12345678" in xml
    assert "0571-87654321" in xml
    assert "lisi@example.com" in xml
    assert "阀门及配件" in xml
    assert "2026年01月01日" in xml
    assert "2026年12月31日" in xml
    assert "常规货期为 【" in xml
    assert "】个工作日" in xml
    assert "3" in xml
    assert "运输方式为" in xml
    assert "快递" in xml
    assert "发货【" in xml
    assert "】小时之前" in xml
    assert "8" in xml
    assert "结算周期：" in xml
    assert "月结30天" in xml
    assert "附件一：采购订单" not in xml
    assert "附件二：年度协议价" not in xml
    assert "priceItem." not in xml
    assert "item." not in xml


def test_render_contract_template_mode_strips_fixed_attachment_section_when_items_empty() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "supplierName": "供应商A",
        "items": [],
    }, config)

    output_path = render_contract(render_data, config, "test_strip_attachment_empty_items", blank_missing=True, table_mode="template")
    try:
        with ZipFile(output_path) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
    finally:
        output_path.unlink(missing_ok=True)

    assert "附件一：设备采购清单" not in xml
    assert xml.count("<w:sectPr") == 1


def test_render_contract_template_mode_keeps_items_rows_without_fixed_attachment_section() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "supplierName": "供应商A",
        "items": [{
            "index": "1",
            "name": "设备A",
            "spec": "A-100",
            "unit": "台",
            "quantity": "2",
            "unitPrice": "100",
            "totalPrice": "200",
            "tagNo": "T-1",
        }],
    }, config)

    output_path = render_contract(render_data, config, "test_strip_attachment_with_items", blank_missing=True, table_mode="template")
    try:
        with ZipFile(output_path) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
    finally:
        output_path.unlink(missing_ok=True)

    assert "设备A" in xml
    assert "附件一：设备采购清单" not in xml


def test_build_attachment_summary_row_for_caigouhetong() -> None:
    config = get_template_config("caigouhetong")
    columns = config.table_bindings["items"]
    row = build_attachment_summary_row(columns, {
        "purchaseSubject": "阀门及配件",
        "totalAmount": "113000",
    })

    assert row == {
        "index": "1",
        "name": "阀门及配件",
        "spec": "详情见附件",
        "unit": "",
        "quantity": "",
        "unitPrice": "",
        "totalPrice": "113000",
        "tagNo": "",
    }


def test_build_attachment_summary_row_for_labor_subcontract() -> None:
    config = get_template_config("laborSubcontract")
    columns = config.table_bindings["items"]
    row = build_attachment_summary_row(columns, {
        "workDescription": "清包工劳务",
        "totalAmount": "50000",
    })

    assert row["index"] == "1"
    assert row["laborItem"] == "清包工劳务"
    assert row["remark"] == "详情见附件"
    assert row["totalPrice"] == "50000"


def test_build_attachment_summary_row_for_professional_subcontract() -> None:
    config = get_template_config("professionalSubcontract")
    columns = config.table_bindings["items"]
    row = build_attachment_summary_row(columns, {
        "engineeringScope": "机电安装",
    })

    assert row["node"] == "机电安装"
    assert row["progressDescription"] == "详情见附件"
    assert row["paymentRate"] == ""


def test_apply_attachment_table_summary_replaces_existing_rows() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "purchaseSubject": "阀门及配件",
        "totalAmount": "113000",
        "items": [{"index": "1", "name": "旧明细", "spec": "旧规格", "unit": "台", "quantity": "1", "unitPrice": "1", "totalPrice": "1", "tagNo": ""}],
    }, config)

    apply_attachment_table_summary(render_data, config)

    assert render_data["items"] == [{
        "index": "1",
        "name": "阀门及配件",
        "spec": "详情见附件",
        "unit": "",
        "quantity": "",
        "unitPrice": "",
        "totalPrice": "113000",
        "tagNo": "",
    }]


def test_render_contract_attachment_mode_inserts_summary_row() -> None:
    config = get_template_config("caigouhetong")
    render_data = merge_render_data({
        "supplierName": "供应商A",
        "purchaseSubject": "阀门及配件",
        "totalAmount": "113000",
        "items": [],
    }, config)

    output_path = render_contract(
        render_data,
        config,
        "test_attachment_summary_row",
        blank_missing=True,
        table_mode="attachment",
    )
    try:
        with ZipFile(output_path) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
    finally:
        output_path.unlink(missing_ok=True)

    assert "阀门及配件" in xml
    assert "详情见附件" in xml
    assert "113000" in xml


def test_render_contract_payment_terms_override_keeps_indent() -> None:
    config = get_template_config("caigouhetong")
    base_data = {
        "contractNo": "HT-001",
        "supplierName": "供应商A",
        "projectName": "项目A",
        "items": [{"index": "1", "name": "阀门", "spec": "DN50", "unit": "台", "quantity": "2", "unitPrice": "10", "totalPrice": "20", "tagNo": "T1"}],
    }
    override_data = merge_render_data({
        **base_data,
        "paymentTermsOverride": "（1）首付款：合同总价的30%。\n（2）到货款：合同总价的70%。",
    }, config)
    output_with_override = render_contract(override_data, config, "test_payment_override_indent")
    try:
        with ZipFile(output_with_override) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
        assert "首付款：合同总价的30%。" in xml
        assert "到货款：合同总价的70%。" in xml
        assert "预付款：合同总价的" not in xml
        override_paragraph = re.search(r"(<w:p[\s\S]*?首付款：合同总价的30%。[\s\S]*?</w:p>)", xml)
        assert override_paragraph
        assert re.search(r'<w:ind [^>]*w:left="559"', override_paragraph.group(1))
    finally:
        output_with_override.unlink(missing_ok=True)

    default_data = merge_render_data(base_data, config)
    output_without_override = render_contract(default_data, config, "test_payment_default_terms")
    try:
        with ZipFile(output_without_override) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
        assert "预付款：合同总价的" in xml
    finally:
        output_without_override.unlink(missing_ok=True)


def test_render_contract_items_content_override_for_supplementary_agreement() -> None:
    config = get_template_config("supplementaryAgreement")
    base_data = {
        "contractNo": "BC-001",
        "supplierName": "供应商A",
        "originalSignYear": "2025",
        "originalSignMonth": "01",
        "originalSignDay": "15",
        "originalContractNo": "HT-001",
        "originalContractTitle": "测试项目设备",
        "amendmentReason": "设计变更。",
        "totalAmount": "1000",
        "totalAmountChinese": "壹仟元整",
        "discountAmountChinese": "玖佰元整",
        "amountWithoutTax": "884.96",
        "deliveryYear": "2025",
        "deliveryMonth": "06",
        "deliveryDay": "30",
        "items": [{
            "index": "1",
            "name": "设备A",
            "spec": "A-100",
            "quantity": "1",
            "unit": "台",
            "unitPrice": "1000",
            "totalPrice": "1000",
            "remark": "",
        }],
    }
    override_data = merge_render_data({
        **base_data,
        "itemsContentOverride": "一、补充内容说明\n二、合计人民币壹仟元整",
    }, config)
    output_with_override = render_contract(override_data, config, "test_supplement_items_override")
    try:
        with ZipFile(output_with_override) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
        assert "补充内容说明" in xml
        assert "合计人民币壹仟元整" in xml
        assert "设备名称" not in xml
        assert "{%tr for item in items %}" not in xml
    finally:
        output_with_override.unlink(missing_ok=True)

    default_data = merge_render_data(base_data, config)
    output_without_override = render_contract(default_data, config, "test_supplement_items_default")
    try:
        with ZipFile(output_without_override) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
        assert "设备名称" in xml
        assert "设备A" in xml
    finally:
        output_without_override.unlink(missing_ok=True)


@pytest.mark.parametrize("template_type", ["caigouhetong", "nonStandardNoInstall", "nonStandardWithInstall"])
def test_equipment_item_rows_render_vertically(template_type: str) -> None:
    config = get_template_config(template_type)
    render_data = merge_render_data({
        "contractNo": f"test-{template_type}",
        "buyerPhone": "0571-00000000",
        "supplierName": "测试供应商",
        "supplierAddress": "测试地址",
        "supplierBank": "测试银行",
        "supplierAccount": "123456",
        "supplierTaxNo": "tax-no",
        "supplierPhone": "13800000000",
        "projectName": "测试项目",
        "purchaseSubject": "测试设备采购",
        "totalAmount": "300",
        "totalAmountChinese": "叁佰元整",
        "taxRate": "13",
        "amountWithoutTax": "265.49",
        "taxAmount": "34.51",
        "items": [
            {
                "index": "1",
                "name": "设备A",
                "spec": "A-100",
                "unit": "台",
                "quantity": "1",
                "unitPrice": "100",
                "totalPrice": "100",
                "tagNo": "TAG-A",
            },
            {
                "index": "2",
                "name": "设备B",
                "spec": "B-200",
                "unit": "台",
                "quantity": "2",
                "unitPrice": "100",
                "totalPrice": "200",
                "tagNo": "TAG-B",
            },
        ],
    }, config)

    rendered_path = render_contract(render_data, config, f"test-{template_type}-vertical-rows", blank_missing=True)
    try:
        with ZipFile(rendered_path) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")
    finally:
        rendered_path.unlink(missing_ok=True)

    assert "{{" not in xml
    assert "{%" not in xml

    row_texts = [
        "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", row))
        for row in re.findall(r"<w:tr[\s\S]*?</w:tr>", xml)
    ]
    row_a = next(index for index, text in enumerate(row_texts) if "设备A" in text)
    row_b = next(index for index, text in enumerate(row_texts) if "设备B" in text)
    assert row_a != row_b
    assert row_a < row_b


def test_contract_generate_passes_confirmed_extracted_data() -> None:
    upload_id = upload_quote(content=b"%PDF-1.4 raw quote")["id"]
    extracted = {"supplierName": "供应商A", "items": []}

    with patch("agent.main.generate_contract") as generate_contract_mock:
        generate_contract_mock.return_value = {
            "contractId": "contract_test",
            "templateType": "caigouhetong",
            "quoteTextLength": 4,
            "dingDrive": {"spaceId": "space1", "fileId": "file1", "fileName": "confirmed.docx"},
        }
        response = client.post(
            "/api/contracts/generate",
            headers=agent_auth_header(),
            json={
                "uploadId": upload_id,
                "templateType": "caigouhetong",
                "quoteText": " 用户确认文本 ",
                "extraInfo": " 补充信息 ",
                "extractedData": extracted,
            },
        )

    assert response.status_code == 200
    generate_contract_mock.assert_called_once_with(upload_id, "caigouhetong", "用户确认文本", "补充信息", extracted, ANY, "auto")
    assert response.json()["download"]["type"] == "agent_proxy"


def test_templates_exist() -> None:
    assert Path("agent/contract/templates/zhanweifu/caigouhetong.docx").exists()
    assert Path("agent/contract/templates/zhanweifu/caigouhetong.placeholders.json").exists()


def test_template_placeholders_match_schema() -> None:
    required_equipment_item_columns = {"index", "name", "spec", "unit", "quantity", "unitPrice", "totalPrice", "tagNo"}
    equipment_templates = {"caigouhetong", "nonStandardNoInstall", "nonStandardWithInstall"}

    for template_type in TEMPLATE_BASENAME:
        config = get_template_config(template_type)
        with ZipFile(template_docx_path(template_type)) as docx:
            xml = docx.read("word/document.xml").decode("utf-8")

        raw_placeholders = set(re.findall(r"\{\{\s*([^}]+?)\s*\}\}", xml))
        placeholders = {re.sub(r"^(?:r|p)\s+", "", placeholder.strip()) for placeholder in raw_placeholders}
        loop_bindings = dict(re.findall(r"\{%\s*(?:tr\s+)?for\s+(\w+)\s+in\s+(\w+)\s*%\}", xml))
        assert placeholders, f"{template_type} should contain docxtpl placeholders"

        scalar_keys = set(config.scalar_keys)
        table_columns = {
            table_name: set(columns)
            for table_name, columns in config.table_bindings.items()
        }
        scalar_placeholders: set[str] = set()
        table_placeholders: dict[str, set[str]] = {table_name: set() for table_name in table_columns}

        for placeholder in placeholders:
            if "." not in placeholder:
                assert placeholder in scalar_keys, f"{template_type}: unknown scalar placeholder {placeholder}"
                scalar_placeholders.add(placeholder)
                continue
            loop_var, column = placeholder.split(".", 1)
            table_name = loop_bindings.get(loop_var)
            assert table_name, f"{template_type}: placeholder {placeholder} is missing a loop binding"
            assert table_name in table_columns, f"{template_type}: unknown table {table_name}"
            assert column in table_columns[table_name], f"{template_type}: unknown table column {placeholder}"
            table_placeholders[table_name].add(column)

        ui_only_scalar_keys = {"deliveryDays"} if template_type in equipment_templates else set()
        assert scalar_keys - ui_only_scalar_keys <= scalar_placeholders, f"{template_type}: schema scalars are not all rendered"
        for table_name, columns in table_columns.items():
            assert columns <= table_placeholders[table_name], f"{template_type}: schema table {table_name} is not fully rendered"

        if template_type in equipment_templates:
            assert required_equipment_item_columns <= table_placeholders["items"]


def test_upload_requires_login_when_enforced(monkeypatch) -> None:
    monkeypatch.setenv("APP_SESSION_SECRET", "enforce-secret-key-123456789012")
    isolated = TestClient(app)
    response = isolated.post(
        "/api/uploads",
        json={
            "originalName": "quote.pdf",
            "mimeType": "application/pdf",
            "size": len(b"x"),
            "data": base64.b64encode(b"x").decode("ascii"),
        },
    )
    assert response.status_code == 401


def test_contract_generate_requires_login_when_enforced(monkeypatch) -> None:
    monkeypatch.setenv("APP_SESSION_SECRET", "enforce-secret-key-123456789012")
    isolated = TestClient(app)
    response = isolated.post(
        "/api/contracts/generate",
        json={
            "uploadId": "upload_missing",
            "templateType": "caigouhetong",
        },
    )
    assert response.status_code == 401


def test_agent_bearer_token_allows_upload_when_enforced(monkeypatch) -> None:
    monkeypatch.setenv("APP_SESSION_SECRET", "enforce-secret-key-123456789012")
    isolated = TestClient(app)
    content = b"%PDF-1.4 hello quote"

    upload = isolated.post(
        "/api/uploads",
        headers=agent_auth_header(),
        json={
            "originalName": "quote.pdf",
            "mimeType": "application/pdf",
            "size": len(content),
            "data": base64.b64encode(content).decode("ascii"),
        },
    )
    assert upload.status_code == 200


